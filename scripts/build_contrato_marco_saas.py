from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Contrato_Marco_SaaS_GE_Control_Plantilla.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
LIGHT = "F5F1F2"
GRID = "D9DEE7"
FONT = "Aptos"


def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def paragraph(doc, text="", bold_lead=None, italic=False, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        font(r2, italic=italic)
    else:
        r = p.add_run(text)
        font(r, italic=italic)
    return p


def clause(doc, number, title, body, page_break_before=False):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    r = p.add_run(f"{number}. {title}")
    font(r, size=14, bold=True, color=WINE)
    for item in body:
        if isinstance(item, tuple):
            lead, text = item
            paragraph(doc, lead + text, bold_lead=lead)
        else:
            paragraph(doc, item)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    font(p.add_run(text))


def variable_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [2300, 4700, 2360]
    headers = ["Variable", "Uso", "Control"]
    for i, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        set_cell_width(cell, widths[i])
        shade(cell, "E9DDE0")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(label), size=9.5, bold=True, color=WINE)
    set_repeat_table_header(table.rows[0])
    for variable, use, control in rows:
        cells = table.add_row().cells
        for i, value in enumerate((variable, use, control)):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            font(p.add_run(value), size=9)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(hp.add_run("GE CONTROL  |  CONTRATO MARCO SaaS"), size=8.5, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("Plantilla privada y versionada  ·  Página "), size=8, color=MUTED)
add_field(fp, "PAGE")
font(fp.add_run(" de "), size=8, color=MUTED)
add_field(fp, "NUMPAGES")

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(92)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("GE CONTROL"), size=14, bold=True, color=WINE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("CONTRATO MARCO DE PRESTACIÓN\nDE SERVICIOS DE SOFTWARE"), size=24, bold=True, color=INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Software como servicio (SaaS) · Portal Transporte"), size=13, color=MUTED)

table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.style = "Table Grid"
meta = [
    ("Versión de plantilla", "{{contrato_version}}"),
    ("Folio", "{{contrato_folio}}"),
    ("Proveedor", "{{proveedor_nombre_legal}} / GE Control"),
    ("Cliente", "{{cliente_nombre_legal}}"),
    ("Fecha", "{{fecha_firma}}"),
]
for row, (label, value) in zip(table.rows, meta):
    set_cell_width(row.cells[0], 2700)
    set_cell_width(row.cells[1], 6660)
    shade(row.cells[0], LIGHT)
    font(row.cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=WINE)
    font(row.cells[1].paragraphs[0].add_run(value), size=9.5)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(34)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("DOCUMENTO PRIVADO · LOS CAMPOS MARCADOS SE COMPLETAN DESDE SUPERADMIN"), size=8.5, bold=True, color=MUTED)
doc.add_page_break()

paragraph(
    doc,
    "CONTRATO MARCO DE PRESTACIÓN DE SERVICIOS DE SOFTWARE COMO SERVICIO que celebran, por una parte, "
    "{{proveedor_nombre_legal}}, con RFC {{proveedor_rfc}}, operando comercialmente como GE Control y con domicilio "
    "en {{proveedor_domicilio_completo}}, representada por {{proveedor_representante}}, en lo sucesivo “GE CONTROL”; "
    "y, por la otra, {{cliente_nombre_legal}}, con RFC {{cliente_rfc}}, domicilio en {{cliente_domicilio_completo}} y "
    "representada por {{cliente_representante}}, en lo sucesivo el “CLIENTE”; conjuntamente, las “PARTES”.",
)

clause(doc, "I", "DECLARACIONES", [
    ("I.1. GE CONTROL. ", "Declara que cuenta con capacidad para contratar y prestar servicios tecnológicos; que utiliza legítimamente el nombre comercial GE Control; y que dispone de la infraestructura, relaciones con proveedores y conocimientos necesarios para prestar el Servicio dentro del alcance contratado."),
    ("I.2. EL CLIENTE. ", "Declara ser una persona moral o persona física con actividad empresarial dedicada o relacionada con el transporte; que los datos entregados son correctos; y que su representante cuenta con facultades suficientes."),
    ("I.3. Ambas PARTES. ", "Reconocen la validez de medios electrónicos, órdenes de servicio, anexos y evidencias digitales, y manifiestan su voluntad de obligarse conforme a este Contrato."),
])

clause(doc, "1", "DEFINICIONES", [
    ("Administrador: ", "usuario nominal autorizado por el CLIENTE para administrar su cuenta y realizar acciones conforme a su rol."),
    ("Carta Porte: ", "CFDI de tipo traslado o ingreso, según corresponda, que incorpora el complemento Carta Porte vigente."),
    ("Cliente contractual: ", "persona que firma este Contrato y puede agrupar una o más suscripciones correspondientes a RFC distintos."),
    ("Datos del Cliente: ", "información, archivos, catálogos, documentos y datos personales ingresados o generados por cuenta del CLIENTE."),
    ("Documento fiscal: ", "archivo XML timbrado y su representación impresa, acuse o documento relacionado."),
    ("Orden de Servicio: ", "instrumento que identifica RFC, plan, precio, vigencia, límites, complementos y condiciones particulares."),
    ("Portal del Operador: ", "complemento opcional que permite al operador consultar asignaciones, documentos y registrar eventos."),
    ("Servicio: ", "acceso remoto a Portal Transporte y a las funciones contratadas bajo modalidad SaaS."),
    ("Viaje fiscal: ", "unidad comercial que permite emitir una Carta Porte y, cuando corresponda y exista previamente aquélla, un CFDI de ingreso relacionado con el mismo servicio."),
])

clause(doc, "2", "OBJETO Y NATURALEZA DEL SERVICIO", [
    "GE CONTROL concede al CLIENTE, durante la vigencia y sujeto al pago, un derecho limitado, no exclusivo, no transferible y no sublicenciable para acceder al Servicio con fines internos de su negocio.",
    "El Servicio es una herramienta tecnológica de apoyo. No constituye asesoría fiscal, contable o jurídica ni sustituye la revisión del CLIENTE y sus asesores. La contratación se realiza para fines empresariales y no para consumo personal, familiar o doméstico.",
])

clause(doc, "3", "ORDEN DE SERVICIO, ACTIVACIÓN Y PRELACIÓN", [
    "Cada RFC operará mediante su propia Orden de Servicio, aun cuando pertenezca al mismo Cliente contractual. Un RFC adicional constituye una suscripción adicional y puede recibir un descuento expresamente documentado.",
    "La activación queda sujeta a validación de identidad, RFC, pago, configuración y documentación requerida. En caso de contradicción prevalecerán: (a) convenio modificatorio firmado; (b) Orden de Servicio para variables comerciales; (c) anexos especializados; y (d) este Contrato.",
])

clause(doc, "4", "CUENTAS, ROLES Y USO AUTORIZADO", [
    "Cada Administrador deberá utilizar una cuenta individual. Los límites de Administradores se fijarán en la Orden de Servicio. Los operadores podrán ser ilimitados dentro del alcance operativo contratado, sin que ello autorice cuentas administrativas compartidas.",
    "El CLIENTE es responsable de altas, bajas, permisos y custodia de credenciales. Deberá notificar accesos indebidos y solicitar inmediatamente la suspensión de usuarios no autorizados.",
    "Queda prohibido revender, sublicenciar, practicar ingeniería inversa, extraer datos de forma abusiva, eludir límites, introducir código malicioso, acceder a otros clientes o utilizar el Servicio para actos ilícitos.",
])

clause(doc, "5", "VIAJES FISCALES Y DOCUMENTOS", [
    "El CFDI de ingreso relacionado con un viaje sólo podrá generarse después de existir una Carta Porte timbrada conforme al flujo habilitado. Un intento sin UUID no consume un viaje fiscal. La obtención de la Carta Porte con UUID inicia el consumo.",
    "La cancelación no repone el viaje. Sustituciones, retimbrados, recuperaciones de duplicados, documentos extraordinarios y errores atribuibles al PAC se regirán por el anexo vigente. No existen timbres ilimitados.",
    "El CLIENTE deberá verificar datos, supuestos fiscales y documentos antes y después del timbrado. GE CONTROL podrá bloquear una emisión que no supere validaciones técnicas, sin asumir la obligación de detectar todo error material o jurídico.",
])

clause(doc, "6", "PRECIOS, IMPUESTOS Y PAGO", [
    "El CLIENTE pagará los importes, impuestos y periodicidad establecidos en cada Orden de Servicio. Los precios podrán actualizarse al renovar, previa comunicación con al menos {{aviso_cambio_precio_dias}} días naturales.",
    "Las suscripciones se pagan por anticipado. Los descuentos deberán indicar porcentaje o monto, causa, vigencia y suscripciones alcanzadas. Ningún descuento se presume permanente.",
    "Ante mora, GE CONTROL enviará avisos razonables. A los diez días naturales podrá suspender nuevas operaciones y timbrados, manteniendo inicialmente consulta y exportación. A los treinta días podrá terminar el Servicio. La suspensión no elimina importes vencidos.",
])

clause(doc, "7", "SOPORTE, DISPONIBILIDAD Y MANTENIMIENTO", [
    "GE CONTROL recibirá solicitudes las veinticuatro horas, todos los días. Los tiempos de primera respuesta y resolución dependen de severidad, información disponible y SLA contratado; la recepción 24/7 no constituye garantía de resolución inmediata.",
    "Podrán realizarse mantenimientos programados, urgentes y actualizaciones de seguridad. Las ventanas, objetivos de disponibilidad, exclusiones y remedios se definen en el Acuerdo de Niveles de Servicio.",
])

clause(doc, "8", "DATOS PERSONALES Y ROLES", [
    "Respecto de prospectos, contactos, cuentas, facturación, seguridad y soporte propios, GE CONTROL actúa como responsable. Respecto de datos de operadores, destinatarios y terceros tratados por instrucciones del CLIENTE, normalmente actúa como persona encargada y el CLIENTE como responsable.",
    "Las PARTES cumplirán el Aviso de Privacidad y el Anexo de Tratamiento y Seguridad. El CLIENTE garantiza que cuenta con bases jurídicas, avisos y autorizaciones necesarias y que sus instrucciones son lícitas.",
], page_break_before=True)

clause(doc, "9", "GEOLOCALIZACIÓN", [
    "Cuando se contrate el Portal del Operador, la funcionalidad actual podrá solicitar ubicación puntual al registrar eventos del viaje, con permiso del dispositivo. No se realizará seguimiento continuo en segundo plano salvo desarrollo futuro expresamente informado y contratado.",
    "El CLIENTE deberá informar a operadores, limitar accesos, evitar vigilancia desproporcionada y cumplir obligaciones laborales y de privacidad. La ubicación se conservará por doce meses, salvo incidencia, reclamación u obligación que justifique ampliar el plazo.",
])

clause(doc, "10", "SEGURIDAD, CONFIDENCIALIDAD E INCIDENTES", [
    "Cada PARTE protegerá la información confidencial con medidas razonables y la usará únicamente para cumplir la relación. Esta obligación subsistirá después de la terminación.",
    "GE CONTROL aplicará controles administrativos, técnicos y físicos proporcionales al riesgo. Notificará incidentes que afecten significativamente derechos o el Servicio conforme a la ley y anexos, sin que una notificación constituya admisión de responsabilidad.",
])

clause(doc, "11", "PROPIEDAD INTELECTUAL Y DATOS", [
    "GE CONTROL conserva todos los derechos sobre software, código, interfaces, documentación, marcas, procesos, configuraciones generales y mejoras. El CLIENTE conserva los derechos que correspondan sobre sus Datos.",
    "El CLIENTE otorga a GE CONTROL autorización limitada para alojar, copiar, transformar y transmitir sus Datos únicamente para prestar, proteger, respaldar y mejorar técnicamente el Servicio, usando información agregada o disociada cuando sea posible.",
])

clause(doc, "12", "PROVEEDORES TECNOLÓGICOS", [
    "El CLIENTE reconoce que el Servicio utiliza proveedores como SW Sapien para certificación, Supabase para datos y autenticación, Render para alojamiento y otros proveedores de correo, monitoreo y conectividad.",
    "GE CONTROL seleccionará y administrará razonablemente a sus proveedores, pero no garantiza continuidad absoluta de servicios externos. Gestionará incidentes y alternativas conforme al SLA y a la viabilidad técnica.",
])

clause(doc, "13", "RESPONSABILIDADES Y GARANTÍAS", [
    "GE CONTROL prestará el Servicio con diligencia razonable y conforme al alcance vigente. No garantiza que toda operación del CLIENTE deba llevar Carta Porte, que los datos proporcionados sean correctos, ni que una aceptación del PAC determine el cumplimiento material de la operación.",
    "El CLIENTE responderá por información, instrucciones, usuarios, legalidad de mercancías, permisos, tratamiento fiscal, relación con operadores y uso de documentos. Deberá mantener copias o exportaciones conforme a su política interna.",
])

clause(doc, "14", "LIMITACIÓN DE RESPONSABILIDAD", [
    "Salvo dolo, culpa grave, violación de confidencialidad, tratamiento ilícito imputable de datos o supuestos que legalmente no puedan limitarse, la responsabilidad acumulada de GE CONTROL derivada de una Orden de Servicio no excederá las cantidades efectivamente pagadas por esa suscripción durante los seis meses anteriores al hecho.",
    "En la medida permitida, ninguna PARTE responderá por daños indirectos, pérdida de oportunidades o lucro cesante no previsible. Esta cláusula no libera a GE CONTROL de entregar el Servicio contratado ni al CLIENTE de pagar importes devengados.",
])

clause(doc, "15", "VIGENCIA, NO RENOVACIÓN Y TERMINACIÓN", [
    "El Contrato inicia en {{fecha_inicio_contrato}} y permanecerá vigente mientras exista una Orden de Servicio, salvo terminación. Cada suscripción tendrá la vigencia indicada en su Orden.",
    "El CLIENTE podrá solicitar no renovación con al menos diez días naturales antes del siguiente periodo. Salvo obligación legal o acuerdo, no habrá reembolso proporcional de periodos iniciados.",
    "Cualquiera podrá terminar por incumplimiento material no subsanado dentro de diez días después de notificación, o inmediatamente ante uso ilícito, riesgo grave de seguridad, insolvencia o reincidencia.",
])

clause(doc, "16", "EXPORTACIÓN, CONSERVACIÓN Y ELIMINACIÓN", [
    "Durante la vigencia el CLIENTE podrá descargar la información disponible. Tras terminación tendrá hasta treinta días para solicitar una exportación razonable, sujeto a pago y viabilidad técnica.",
    "Política base: documentos fiscales, al menos cinco años; evidencias operativas, veinticuatro meses; geolocalización, prospectos y registros técnicos, doce meses. Los plazos pueden ampliarse por obligaciones, incidencias o defensa de responsabilidades.",
    "Concluidos los plazos se eliminarán o anonimizarán datos, previo bloqueo cuando corresponda. La terminación no obliga a destruir información cuya conservación sea legalmente necesaria.",
])

clause(doc, "17", "NOTIFICACIONES", [
    "Las comunicaciones contractuales se enviarán a {{proveedor_correo_contratos}} y {{cliente_correo_contratos}}. Privacidad: privacidad@gecontrol.mx. Un cambio de correo surtirá efectos al ser confirmado.",
    "Las notificaciones electrónicas se considerarán recibidas cuando exista evidencia de entrega, acceso, aceptación o respuesta, sin perjuicio de formalidades legales específicas.",
])

clause(doc, "18", "FUERZA MAYOR, CESIÓN Y SUBCONTRATACIÓN", [
    "Ninguna PARTE será responsable por incumplimientos causados por eventos fuera de su control razonable mientras mitigue y comunique. Las obligaciones de pago ya causadas no se extinguen.",
    "El CLIENTE no cederá el Contrato sin consentimiento. GE CONTROL podrá utilizar encargados tecnológicos y ceder a una entidad sucesora de su negocio, asegurando continuidad de obligaciones y avisando cuando corresponda.",
])

clause(doc, "19", "MODIFICACIONES, INTEGRIDAD Y FIRMA ELECTRÓNICA", [
    "Este Contrato, órdenes y anexos integran el acuerdo. Las modificaciones requieren versión identificable y aceptación válida. Una publicación web no modificará retroactivamente un documento firmado.",
    "Las PARTES reconocen firmas autógrafas, electrónicas, códigos de aceptación y evidencias técnicas conforme a su fuerza probatoria aplicable. Cada PDF emitido conservará folio, versión, fecha, variables, hash y auditoría.",
], page_break_before=True)

clause(doc, "20", "LEY Y CONTROVERSIAS", [
    "Las PARTES procurarán negociar de buena fe durante treinta días. Este Contrato se regirá por las leyes mexicanas y, salvo pacto o competencia imperativa, las PARTES se someten a los tribunales competentes de {{jurisdiccion_estado}}, renunciando al fuero que pudiera corresponderles por domicilio futuro.",
])

doc.add_page_break()
p = doc.add_paragraph(style="Heading 1")
font(p.add_run("FIRMAS"), size=16, bold=True, color=WINE)
paragraph(doc, "Leído el presente Contrato y enteradas las PARTES de su alcance, lo aceptan en la fecha indicada.")
sig = doc.add_table(rows=4, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
sig.autofit = False
for row in sig.rows:
    set_cell_width(row.cells[0], 4680)
    set_cell_width(row.cells[1], 4680)
labels = [
    ("GE CONTROL", "EL CLIENTE"),
    ("{{proveedor_nombre_legal}}", "{{cliente_nombre_legal}}"),
    ("{{proveedor_representante}}", "{{cliente_representante}}"),
    ("Firma: __________________________", "Firma: __________________________"),
]
for row, pair in zip(sig.rows, labels):
    for cell, value in zip(row.cells, pair):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        font(p.add_run(value), size=9.5, bold=row is sig.rows[0])

doc.add_page_break()
p = doc.add_paragraph(style="Heading 1")
font(p.add_run("ANEXO TÉCNICO DE PLANTILLA · CAMPOS EDITABLES"), size=16, bold=True, color=WINE)
paragraph(doc, "Este anexo es de control interno para Superadmin y no necesariamente se adjunta al contrato firmado.", italic=True)

variable_table(doc, [
    ("{{proveedor_*}}", "Identidad, RFC, domicilio, representante y correos.", "Perfil legal versionado"),
    ("{{cliente_*}}", "Identidad, RFC, domicilio, representante y contactos.", "Expediente contractual"),
    ("{{contrato_*}}", "Folio, versión, fechas, jurisdicción y estado.", "Inmutable al emitir"),
    ("{{orden_servicio_*}}", "Plan, RFC, precio, descuento, límites y complemento.", "Editable por suscripción"),
    ("{{aviso_cambio_precio_dias}}", "Anticipación para cambio al renovar.", "Política versionada"),
    ("{{jurisdiccion_estado}}", "Entidad federativa pactada.", "Revisión jurídica"),
])

p = doc.add_paragraph(style="Heading 2")
font(p.add_run("Cláusulas protegidas"), size=12, bold=True, color=WINE)
for text in [
    "Definiciones legales y fiscales.",
    "Responsabilidad sobre datos y documentos.",
    "Protección de datos, geolocalización y seguridad.",
    "Propiedad intelectual y confidencialidad.",
    "Limitación de responsabilidad, terminación y controversias.",
]:
    bullet(doc, text)

p = doc.add_paragraph(style="Heading 2")
font(p.add_run("Reglas de versionado"), size=12, bold=True, color=WINE)
for text in [
    "Un borrador puede editarse; un documento enviado o firmado es inmutable.",
    "Todo cambio comercial genera nueva versión de Orden de Servicio.",
    "Todo cambio legal genera nueva versión de plantilla y revisión.",
    "El PDF debe conservar los valores resueltos, no sólo los nombres de variables.",
    "Guardar hash, autor, fecha, destinatarios, evidencia de envío y aceptación.",
]:
    bullet(doc, text)

doc.save(DEST)
print(DEST)
