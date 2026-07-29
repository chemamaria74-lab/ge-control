from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Anexo_Tratamiento_Seguridad_Datos_GE_Control.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.3, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def subheading(doc, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(title), size=11.5, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=9.8)


def kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_width(cells[0], 2800)
        set_width(cells[1], 6560)
        shade(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.2, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.2)
    return table


def grid(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        set_width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8.8, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(value), size=font_size)
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.3)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.14
for style_name, size in (("Heading 1", 14), ("Heading 2", 11.5)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(11)
    style.paragraph_format.space_after = Pt(5)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  TRATAMIENTO Y SEGURIDAD DE DATOS"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Anexo de Tratamiento y Seguridad v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("ANEXO DE TRATAMIENTO Y SEGURIDAD DE DATOS"), size=21, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(15)
set_font(subtitle.add_run("Portal Transporte · Anexo del Contrato Marco SaaS"), size=12.5, color=MUTED)

kv_table(doc, [
    ("Versión del Anexo", "{{dpa_version}}"),
    ("Vigencia", "{{dpa_fecha_vigencia}}"),
    ("Contrato Marco", "{{contrato_marco_folio}}"),
    ("Orden de Servicio", "{{orden_servicio_folio}}"),
    ("Cliente / Responsable", "{{cliente_nombre_legal}}"),
    ("RFC de la suscripción", "{{suscripcion_rfc}}"),
    ("Contacto del Cliente", "{{cliente_contacto_privacidad}}"),
])

para(doc, "Este Anexo forma parte integral del Contrato Marco de Prestación de Servicios SaaS celebrado entre el CLIENTE y María José Mejía Ornelas, quien opera comercialmente como GE Control. Regula el tratamiento que GE Control realiza por cuenta del CLIENTE en Portal Transporte y sus complementos habilitados.")

heading(doc, "1", "OBJETO, PRELACIÓN Y MARCO JURÍDICO")
para(doc, "El objeto es documentar instrucciones, responsabilidades, medidas de seguridad, subcontrataciones y mecanismos de asistencia respecto de Datos Personales del Cliente. En caso de contradicción sobre protección de datos, prevalecerá este Anexo; para materias comerciales prevalecerán el Contrato Marco y la Orden de Servicio.")
para(doc, "Las Partes cumplirán la Ley Federal de Protección de Datos Personales en Posesión de los Particulares publicada el 20 de marzo de 2025, su reforma vigente, su Reglamento y disposiciones aplicables. Las referencias a una autoridad significan la Secretaría Anticorrupción y Buen Gobierno o la autoridad que legalmente la sustituya.")

heading(doc, "2", "DEFINICIONES")
grid(doc, ["Término", "Significado"], [
    ("Datos del Cliente", "Información, archivos y datos que el CLIENTE o sus usuarios incorporan, generan o consultan en el Servicio."),
    ("Datos Personales", "Información concerniente a una persona física identificada o identificable."),
    ("Responsable", "El CLIENTE, cuando decide finalidades y medios esenciales del tratamiento."),
    ("Encargado", "GE Control, cuando trata Datos Personales por cuenta y conforme a instrucciones del CLIENTE."),
    ("Subencargado", "Tercero contratado por GE Control para apoyar el tratamiento por cuenta del CLIENTE."),
    ("Incidente", "Pérdida, destrucción, alteración, acceso, adquisición, uso o divulgación no autorizada de Datos Personales."),
    ("Servicios", "Portal Transporte, Portal del Operador y funciones habilitadas en la Orden de Servicio."),
], [2200, 7160])

heading(doc, "3", "ROLES Y TRATAMIENTOS PROPIOS")
para(doc, "El CLIENTE es Responsable de los datos de operadores, conductores, remitentes, destinatarios, contactos, empleados y terceros que carga o genera en el Servicio. GE Control actúa como Encargado respecto de esos datos.")
para(doc, "GE Control actúa separadamente como Responsable respecto de datos indispensables para administrar la relación comercial, facturar, autenticar cuentas, proteger su infraestructura, prevenir abuso, atender soporte, conservar auditoría propia y cumplir obligaciones legales. Dichos tratamientos se rigen por el Aviso de Privacidad Integral y no por instrucciones del CLIENTE.")

heading(doc, "4", "INSTRUCCIONES DOCUMENTADAS")
para(doc, "El CLIENTE instruye a GE Control a recibir, alojar, organizar, consultar, relacionar, respaldar, transmitir, recuperar, exportar, bloquear y eliminar Datos Personales únicamente para prestar, proteger y soportar los Servicios contratados.")
para(doc, "Configurar módulos, cargar datos, solicitar soporte, utilizar APIs, autorizar usuarios o emitir instrucciones mediante contactos autorizados constituyen instrucciones documentadas. GE Control no venderá los Datos del Cliente ni los utilizará para publicidad propia.")
para(doc, "Si GE Control considera que una instrucción infringe la legislación, lo informará y podrá suspender su ejecución hasta que sea aclarada. No estará obligado a ejecutar instrucciones técnicamente inviables, inseguras, ilícitas o ajenas al Servicio contratado.")

heading(doc, "5", "OBLIGACIONES DEL CLIENTE")
for item in [
    "Contar con base jurídica y, cuando corresponda, consentimiento para recabar y tratar los datos.",
    "Entregar avisos de privacidad adecuados a operadores, empleados y demás titulares antes del tratamiento.",
    "Limitar los datos a los necesarios y evitar datos sensibles, salvo necesidad lícita y protección reforzada.",
    "Mantener exactos sus datos, configurar permisos, retirar accesos oportunamente y proteger credenciales.",
    "Informar a operadores sobre la geolocalización por evento, su finalidad, conservación y destinatarios.",
    "Atender derechos ARCO y decisiones laborales, fiscales u operativas; GE Control sólo prestará asistencia.",
    "No utilizar el Servicio para vigilancia secreta, discriminación, hostigamiento ni finalidades incompatibles.",
]:
    bullet(doc, item)

heading(doc, "6", "PERSONAL, CONFIDENCIALIDAD Y ACCESO")
para(doc, "GE Control limitará el acceso a personal o colaboradores que lo requieran para sus funciones, bajo deber de confidencialidad durante y después de su participación. Aplicará el principio de mínimo privilegio, mecanismos de autenticación y revocación de acceso.")
para(doc, "El acceso de soporte a Datos del Cliente deberá responder a un ticket, incidente, mantenimiento autorizado, obligación legal o necesidad documentada. GE Control conservará trazabilidad razonable de accesos administrativos conforme a sus capacidades técnicas.")

heading(doc, "7", "MEDIDAS DE SEGURIDAD", page_break_before=True)
para(doc, "GE Control implementará y mantendrá medidas administrativas, técnicas y físicas proporcionales al volumen, contexto, sensibilidad, amenazas conocidas y consecuencias posibles. Las medidas no constituyen garantía de seguridad absoluta ni certificación ISO, SOC o equivalente.")
grid(doc, ["Dominio", "Medidas comprometidas"], [
    ("Gobierno", "Responsables definidos, confidencialidad, revisión de riesgos, gestión de cambios y respuesta a incidentes."),
    ("Identidad y acceso", "Cuentas individuales, roles, mínimo privilegio, expiración o revocación de sesiones y controles de acceso administrativo."),
    ("Separación lógica", "Validación del contexto de cliente y RFC, filtros autoritativos y controles de membresía en componentes aplicables."),
    ("Transmisión y almacenamiento", "TLS en tránsito en servicios compatibles; almacenamiento privado y protecciones provistas por la infraestructura contratada."),
    ("Desarrollo", "Control de cambios, revisión proporcional al riesgo, manejo reservado de secretos y pruebas antes de liberaciones materiales."),
    ("Registro y monitoreo", "Registros de autenticación, errores y eventos relevantes; monitoreo y conservación conforme a la política aplicable."),
    ("Continuidad", "Respaldos y procedimientos de recuperación sujetos a los objetivos RPO/RTO previstos en el SLA."),
    ("Eliminación", "Borrado lógico, bloqueo, anonimización o eliminación según funcionalidad, conservación contractual y respaldos rotativos."),
], [2300, 7060], 8.2)
para(doc, "Los controles pueden evolucionar para responder a riesgos y tecnología, siempre que el nivel general de protección no disminuya materialmente. El CLIENTE reconoce que la seguridad también depende de sus usuarios, dispositivos, redes y configuraciones.")

heading(doc, "8", "SUBENCARGADOS Y SERVICIOS DE TERCEROS")
para(doc, "El CLIENTE otorga autorización general para utilizar los Subencargados identificados en el Apéndice C. GE Control exigirá obligaciones de protección acordes con la naturaleza del servicio y seguirá siendo responsable contractualmente de sus propias obligaciones frente al CLIENTE.")
para(doc, "GE Control comunicará cambios materiales en la lista por correo, plataforma u otro medio registrado, cuando sea razonablemente posible con al menos quince días naturales de anticipación. El CLIENTE podrá formular una objeción fundada en protección de datos dentro de ese plazo. Las Partes buscarán una alternativa razonable; si no existe, podrán terminar únicamente el servicio afectado sin penalidad futura.")

heading(doc, "9", "REMISIONES, UBICACIÓN Y ACCESO TRANSFRONTERIZO")
para(doc, "Las comunicaciones entre el Responsable, GE Control y sus Subencargados para prestar el Servicio se considerarán remisiones cuando legalmente corresponda. Algunos proveedores pueden almacenar o acceder a información desde infraestructura fuera de México.")
para(doc, "GE Control procurará que los destinatarios queden sujetos a obligaciones de confidencialidad, seguridad, finalidad y supresión. El CLIENTE autoriza dichas remisiones para ejecutar el Contrato, sin perjuicio de las obligaciones de información que le correspondan frente a titulares.")

heading(doc, "10", "DERECHOS DE TITULARES Y ASISTENCIA")
para(doc, "El CLIENTE es responsable de recibir, verificar y resolver solicitudes de acceso, rectificación, cancelación, oposición, revocación o limitación relacionadas con Datos del Cliente. Si GE Control recibe directamente una solicitud identificable, la remitirá al contacto registrado, salvo prohibición legal.")
para(doc, "Considerando la naturaleza del tratamiento, GE Control proporcionará asistencia razonable mediante funcionalidades disponibles, búsquedas, exportaciones, correcciones o eliminación. Trabajo extraordinario, recuperación histórica o desarrollo especial podrá cotizarse previamente.")

heading(doc, "11", "INCIDENTES DE SEGURIDAD", page_break_before=True)
para(doc, "GE Control notificará al contacto de seguridad del CLIENTE sin dilación indebida y, en la medida razonablemente posible, dentro de las veinticuatro horas siguientes a confirmar un Incidente que afecte Datos del Cliente. Una alerta no confirmada o intento bloqueado no activa por sí solo esta obligación.")
para(doc, "La notificación incluirá progresivamente, según se conozca: naturaleza y alcance; categorías de datos y titulares; fecha estimada; consecuencias probables; medidas de contención y recuperación; recomendaciones; y contacto de seguimiento.")
para(doc, "GE Control investigará, contendrá y documentará el Incidente; preservará evidencia razonable y cooperará. El CLIENTE decidirá y realizará las notificaciones a titulares o autoridades como Responsable, con apoyo razonable de GE Control. Ninguna notificación constituye admisión de responsabilidad.")

heading(doc, "12", "CONSERVACIÓN, DEVOLUCIÓN Y ELIMINACIÓN")
para(doc, "Durante la vigencia, el CLIENTE podrá exportar información mediante las funciones disponibles. Al terminar la suscripción, GE Control mantendrá una ventana de exportación conforme al Contrato y a la Política de Soporte, Respaldos, Conservación y Exportación.")
para(doc, "Vencida la ventana, GE Control eliminará o anonimizará los Datos del Cliente en sistemas activos, salvo conservación exigida por ley, defensa de derechos, bloqueo o instrucción escrita válida. Las copias en respaldos rotativos se eliminarán por ciclo y permanecerán aisladas de uso ordinario.")
grid(doc, ["Categoría", "Plazo operativo base"], [
    ("Documentos fiscales y trazabilidad", "Al menos 5 años o el plazo legal mayor."),
    ("Evidencias operativas", "24 meses."),
    ("Geolocalización asociada a eventos", "12 meses, salvo incidencia o reclamación."),
    ("Registros de seguridad", "12 meses."),
], [6000, 3360])

heading(doc, "13", "AUDITORÍA Y EVIDENCIA DE CUMPLIMIENTO")
para(doc, "A solicitud escrita no más de una vez por año, GE Control proporcionará información razonable para acreditar cumplimiento, sujeta a confidencialidad y sin revelar secretos, datos de otros clientes o vulnerabilidades explotables.")
para(doc, "Una auditoría adicional sólo procederá ante Incidente confirmado, requerimiento de autoridad o indicio objetivo de incumplimiento material. Se realizará en horario hábil, con alcance acordado, sin afectar el Servicio y a costo del CLIENTE, salvo incumplimiento material imputable a GE Control.")

heading(doc, "14", "REQUERIMIENTOS DE AUTORIDAD")
para(doc, "GE Control podrá divulgar datos cuando lo exija una orden o disposición válida. Salvo prohibición, notificará al CLIENTE antes de responder y limitará la entrega a lo legalmente requerido. Las solicitudes dirigidas al CLIENTE podrán remitirse para su atención.")

heading(doc, "15", "RESPONSABILIDAD")
para(doc, "Cada Parte responde por sus propias obligaciones y decisiones. El CLIENTE responde por licitud, avisos, consentimiento, calidad de los datos, instrucciones, usuarios y uso laboral o fiscal. GE Control responde por ejecutar las instrucciones aceptadas y mantener las medidas comprometidas.")
para(doc, "Los límites, exclusiones y procedimiento de reclamación del Contrato Marco aplican a este Anexo, salvo que una norma imperativa disponga otra cosa. Nada limita derechos de las personas titulares ni facultades de la autoridad.")

heading(doc, "16", "VIGENCIA Y CAMBIOS")
para(doc, "Este Anexo inicia con la Orden de Servicio y permanece mientras GE Control trate Datos del Cliente. Sus obligaciones de confidencialidad, devolución, eliminación, cooperación y responsabilidad subsisten durante los plazos aplicables.")
para(doc, "Los cambios materiales se versionarán y no modificarán retroactivamente obligaciones respecto de tratamientos concluidos. Cada documento aceptado conservará versión, fecha, cliente, RFC y evidencia de aceptación.")

heading(doc, "APÉNDICE A", "DETALLES DEL TRATAMIENTO", page_break_before=True)
grid(doc, ["Elemento", "Descripción"], [
    ("Objeto", "Prestación, seguridad, soporte y continuidad de Portal Transporte y complementos contratados."),
    ("Duración", "Vigencia de la suscripción más conservación, bloqueo, exportación y eliminación aplicables."),
    ("Operaciones", "Recepción, registro, organización, relación, consulta, transmisión, respaldo, recuperación, exportación y eliminación."),
    ("Titulares", "Operadores, conductores, remitentes, destinatarios, empleados, contactos, representantes y terceros incorporados por el CLIENTE."),
    ("Datos", "Identificación, contacto, laborales, licencia, vehículo, viaje, mercancía, evidencia, documentos fiscales, acceso, seguridad y geolocalización por evento."),
    ("Datos sensibles", "No requeridos como regla general. Su carga exige necesidad lícita, minimización y salvaguardas del CLIENTE."),
    ("Frecuencia", "Continua durante el uso del Servicio; geolocalización sólo en eventos habilitados y con permiso del dispositivo."),
], [2200, 7160])

heading(doc, "APÉNDICE B", "CONTACTOS Y PARÁMETROS")
kv_table(doc, [
    ("Privacidad GE Control", "privacidad@gecontrol.mx"),
    ("Seguridad / incidentes GE Control", "{{gecontrol_contacto_seguridad}}"),
    ("Privacidad del Cliente", "{{cliente_contacto_privacidad}}"),
    ("Seguridad / incidentes del Cliente", "{{cliente_contacto_seguridad}}"),
    ("Ubicación principal informada", "{{region_principal_datos}}"),
    ("Ventana de exportación al terminar", "{{dias_ventana_exportacion}} días naturales"),
])

heading(doc, "APÉNDICE C", "SUBENCARGADOS AUTORIZADOS")
grid(doc, ["Proveedor", "Finalidad", "Datos potenciales", "Ubicación"], [
    ("Supabase", "Base de datos, almacenamiento y autenticación.", "Cuenta, contenido, documentos, registros y metadatos.", "{{region_supabase}}"),
    ("Render", "Alojamiento y ejecución de la aplicación.", "Solicitudes, contenido procesado, IP y registros técnicos.", "{{region_render}}"),
    ("SW Sapien", "Validación, certificación y operación fiscal.", "RFC, CFDI, Carta Porte, XML y datos fiscales relacionados.", "México / según servicio contratado"),
    ("{{proveedor_correo}}", "Envío de avisos, soporte y comunicaciones transaccionales.", "Nombre, correo, asunto y contenido necesario.", "{{region_correo}}"),
], [1600, 2600, 3560, 1600], 7.8)
para(doc, "La lista vigente deberá conservar fecha de actualización y no podrá alterarse dentro de un documento ya aceptado. Los cambios futuros se registrarán como una nueva versión.", italic=True)

heading(doc, "APÉNDICE D", "ACEPTACIÓN")
para(doc, "Las Partes manifiestan contar con facultades suficientes y aceptan el presente Anexo. La firma electrónica, aceptación verificable o firma autógrafa producirá los efectos previstos en el Contrato Marco.")
grid(doc, ["POR EL CLIENTE", "POR GE CONTROL"], [
    ("Nombre: {{cliente_representante}}\nCargo: {{cliente_cargo}}\nFirma: __________________________\nFecha: {{cliente_fecha_firma}}",
     "María José Mejía Ornelas\nGE Control\nFirma: __________________________\nFecha: {{gecontrol_fecha_firma}}"),
], [4680, 4680], 9.0)

heading(doc, "APÉNDICE INTERNO", "CAMPOS PARA SUPERADMIN")
para(doc, "Este apéndice es operativo y deberá omitirse del PDF firmado. Las cláusulas legales se conservan versionadas; sólo los campos autorizados se completan desde Superadmin.", italic=True)
grid(doc, ["Grupo", "Campos editables"], [
    ("Identificación", "dpa_version, fecha_vigencia, contrato_marco_folio, orden_servicio_folio, cliente, RFC."),
    ("Contactos", "Privacidad y seguridad de ambas Partes."),
    ("Infraestructura", "Regiones, proveedor de correo y lista versionada de Subencargados."),
    ("Terminación", "Días de exportación, siempre dentro de la política aprobada."),
    ("Firma", "Representante, cargo, fecha, método, IP/evidencia y hash del PDF."),
], [2200, 7160])

doc.core_properties.title = "Anexo de Tratamiento y Seguridad de Datos - GE Control"
doc.core_properties.subject = "Portal Transporte - Acuerdo responsable-encargado"
doc.core_properties.author = "GE Control"
doc.core_properties.keywords = "protección de datos, encargado, SaaS, seguridad, Portal Transporte"
doc.save(DEST)
print(DEST)
