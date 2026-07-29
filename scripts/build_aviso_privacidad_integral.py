from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Aviso_Privacidad_Integral_GE_Control.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=10)


def grid(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=9, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(value), size=8.7)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  AVISO DE PRIVACIDAD INTEGRAL"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Aviso de Privacidad Integral v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(20)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("AVISO DE PRIVACIDAD INTEGRAL"), size=23, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(3)
set_font(subtitle.add_run("GE Control · Portal Transporte"), size=13, color=MUTED)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(16)
set_font(meta.add_run("Versión 1.0 · Vigente desde el 28 de julio de 2026"), size=9, color=MUTED)

para(doc, "Este Aviso informa las características principales del tratamiento de datos personales realizado en el sitio, Superadmin, Portal Transporte, Portal del Operador, procesos comerciales, soporte y documentos relacionados.")

heading(doc, "1", "IDENTIDAD Y DOMICILIO DE LA RESPONSABLE")
para(doc, "María José Mejía Ornelas, operando comercialmente como GE Control, con domicilio para oír y recibir notificaciones en Carretera Sierra del Humo 211, colonia Bosques del Prado Sur, código postal 20130, Aguascalientes, Aguascalientes, México, es responsable del tratamiento que realiza para finalidades comerciales, contractuales, facturación, soporte, seguridad y administración del Servicio.")
para(doc, "La persona o área encargada de datos personales recibirá solicitudes en privacidad@gecontrol.mx.")

heading(doc, "2", "ROLES EN EL TRATAMIENTO")
para(doc, "GE Control actúa como responsable respecto de prospectos, contactos, clientes, administradores, facturación, seguridad, soporte y operación propia. Cuando trata datos de operadores, destinatarios, remitentes u otros terceros exclusivamente por instrucciones de un cliente empresarial, GE Control actúa normalmente como persona encargada y el cliente como responsable.")
para(doc, "El cliente debe proporcionar sus propios avisos, contar con base jurídica e impartir instrucciones lícitas. La relación responsable-encargado se regula también mediante el Anexo de Tratamiento y Seguridad.")

heading(doc, "3", "DATOS PERSONALES TRATADOS")
grid(doc, ["Categoría", "Ejemplos"], [
    ("Identificación", "Nombre, firma, RFC, CURP cuando resulte necesaria, puesto y empresa."),
    ("Contacto", "Correo, teléfono, domicilio y medios para notificaciones."),
    ("Laborales y profesionales", "Puesto, empresa, relación con el cliente, licencia y datos del operador."),
    ("Fiscales y facturación", "RFC, régimen, código postal, documentos, folios, UUID y datos de cobro."),
    ("Operativos", "Viajes, rutas, vehículos, mercancías, asignaciones, incidencias, evidencias y liquidaciones."),
    ("Acceso y seguridad", "Usuario, rol, IP, dispositivo, sesión, eventos, intentos, auditoría y errores."),
    ("Ubicación", "Coordenadas, precisión, fecha, hora, viaje y evento cuando el dispositivo lo autoriza."),
    ("Documentales", "Archivos, fotografías, XML, PDF y evidencias cargadas por usuarios autorizados."),
], [2500, 6860])

para(doc, "GE Control no solicita datos personales sensibles como regla general. El cliente debe evitar cargarlos. Si un documento contiene datos sensibles, deberá contar con la base jurídica y consentimiento expreso cuando resulte exigible.")

heading(doc, "4", "FINALIDADES PRIMARIAS")
for item in [
    "Identificar, autenticar y administrar usuarios, roles, clientes, RFC y suscripciones.",
    "Prestar, configurar, mantener, monitorear y proteger Portal Transporte y sus complementos.",
    "Crear y administrar viajes, Carta Porte, CFDI de ingreso, documentos, evidencias y expedientes.",
    "Asignar operaciones, consultar eventos y proporcionar el Portal del Operador.",
    "Atender soporte, investigar incidentes, prevenir fraude y conservar trazabilidad.",
    "Gestionar prospectos, demostraciones solicitadas, cotizaciones, contratos, cobros y renovaciones.",
    "Emitir comprobantes, cumplir obligaciones contractuales, fiscales, administrativas y requerimientos de autoridad.",
    "Respaldar, recuperar, exportar, bloquear, conservar y eliminar información conforme a los plazos aplicables.",
]:
    bullet(doc, item)

heading(doc, "5", "FINALIDADES SECUNDARIAS")
para(doc, "Con consentimiento o base jurídica aplicable, los datos de contacto podrán utilizarse para encuestas, novedades, invitaciones, promociones y seguimiento comercial no indispensable. La persona puede oponerse o solicitar que cesen en privacidad@gecontrol.mx, sin afectar el Servicio contratado.")

heading(doc, "6", "GEOLOCALIZACIÓN")
para(doc, "Actualmente, el Portal del Operador solicita ubicación precisa únicamente cuando la persona registra determinados eventos del viaje y el dispositivo concede permiso. No realiza seguimiento continuo en segundo plano.")
para(doc, "La ubicación permite documentar eventos como llegada, carga, tránsito, incidencia o entrega; puede ser consultada por personal autorizado del cliente y por GE Control cuando sea necesario para soporte, seguridad o cumplimiento. El cliente deberá informar a sus operadores, limitar accesos y evitar vigilancia desproporcionada o ajena a la operación.")

heading(doc, "7", "TECNOLOGÍAS DE RECOLECCIÓN AUTOMÁTICA")
para(doc, "El sitio y la plataforma pueden utilizar cookies técnicas, almacenamiento local, tokens de sesión y registros para autenticación, preferencias, seguridad y continuidad. Estas tecnologías pueden recabar IP, navegador, dispositivo, fecha, hora, ruta consultada y eventos técnicos.")
para(doc, "Las tecnologías estrictamente necesarias no pueden deshabilitarse sin afectar el funcionamiento. Si se incorporan tecnologías publicitarias o analíticas no necesarias, se informará y habilitará el mecanismo de consentimiento correspondiente.")

heading(doc, "8", "PERSONAS ENCARGADAS Y PROVEEDORES")
para(doc, "GE Control utiliza proveedores que pueden tratar datos por cuenta de la responsable o del cliente, entre ellos Supabase para datos y autenticación, Render para alojamiento, SW Sapien para certificación fiscal y proveedores de correo, respaldo, monitoreo y conectividad.")
para(doc, "Se limitará el tratamiento a lo necesario y se procurarán obligaciones de confidencialidad, seguridad, supresión y subcontratación acordes con el servicio. Algunos proveedores pueden operar infraestructura fuera de México.")

heading(doc, "9", "TRANSFERENCIAS")
para(doc, "Los datos podrán comunicarse sin consentimiento adicional cuando la ley lo permita, incluyendo requerimientos de autoridad, cumplimiento de una relación jurídica, protección de derechos o comunicación al cliente responsable que instruyó el tratamiento.")
para(doc, "Si GE Control pretende realizar una transferencia distinta que requiera consentimiento, informará identidad o categoría del receptor, finalidad y medio para aceptar o rechazar. GE Control no vende datos personales.")

heading(doc, "10", "CONSERVACIÓN, BLOQUEO Y ELIMINACIÓN")
grid(doc, ["Información", "Plazo operativo base"], [
    ("Documentos fiscales y trazabilidad", "Al menos 5 años."),
    ("Evidencias y documentos operativos", "24 meses."),
    ("Ubicación asociada a eventos", "12 meses, salvo incidencia o reclamación."),
    ("Prospectos que no contrataron", "12 meses desde la última interacción."),
    ("Registros técnicos y de seguridad", "12 meses."),
], [5600, 3760])
para(doc, "Un plazo podrá ampliarse por disposición jurídica, contrato, procedimiento, investigación, ejercicio de derechos o defensa de responsabilidades. Concluida la finalidad, los datos serán bloqueados cuando corresponda y posteriormente eliminados o anonimizados.")

heading(doc, "11", "SEGURIDAD Y VULNERACIONES")
para(doc, "GE Control mantiene medidas administrativas, técnicas y físicas razonables según riesgo, incluyendo controles de acceso, separación por cliente y RFC, almacenamiento privado, auditoría, respaldos y gestión de incidentes.")
para(doc, "Cuando una vulneración afecte significativamente derechos patrimoniales o morales, se notificará a las personas afectadas conforme a la legislación aplicable, indicando la naturaleza, datos comprometidos, recomendaciones, acciones correctivas y medio de contacto cuando la investigación lo permita.")

heading(doc, "12", "DERECHOS ARCO")
para(doc, "La persona titular o su representante puede solicitar acceso, rectificación, cancelación u oposición respecto de sus datos mediante privacidad@gecontrol.mx.")
para(doc, "La solicitud deberá incluir: nombre y medio para recibir respuesta; documentos que acrediten identidad o representación; descripción clara de los datos; derecho que desea ejercer; y elementos que faciliten su localización.")
para(doc, "GE Control comunicará su determinación dentro de veinte días contados desde la recepción de una solicitud completa y, si resulta procedente, la hará efectiva dentro de los quince días siguientes. Los plazos podrán ampliarse una sola vez por causa justificada.")

heading(doc, "13", "REVOCACIÓN Y LIMITACIÓN")
para(doc, "Cuando el tratamiento dependa del consentimiento, podrá solicitarse su revocación. También puede solicitarse limitar el uso o divulgación. La revocación no tendrá efectos retroactivos y puede no proceder respecto de tratamientos necesarios para un contrato, obligación legal, seguridad o defensa de responsabilidades.")
para(doc, "Los permisos de ubicación pueden administrarse desde el dispositivo. Desactivarlos puede impedir registrar eventos que requieran evidencia, pero no autoriza capturas ocultas.")

heading(doc, "14", "DECISIONES AUTOMATIZADAS")
para(doc, "GE Control puede aplicar reglas automáticas para validaciones, seguridad, límites, alertas y bloqueo técnico. Estas reglas no están destinadas a evaluar sin intervención humana el rendimiento laboral del operador ni a producir decisiones jurídicas sobre su relación de trabajo. Las decisiones comerciales o disciplinarias corresponden al cliente.")

heading(doc, "15", "CAMBIOS AL AVISO")
para(doc, "Los cambios se publicarán en gecontrol.mx o en la dirección habilitada dentro de la plataforma, identificando versión y fecha. Los cambios materiales se comunicarán mediante plataforma, correo u otro medio adecuado. Las aceptaciones conservarán versión, fecha, usuario y evidencia técnica disponible.")

heading(doc, "16", "AUTORIDAD Y CONTACTO")
para(doc, "La persona titular puede acudir ante la Secretaría Anticorrupción y Buen Gobierno o la autoridad competente en materia de protección de datos personales si considera vulnerados sus derechos.")
para(doc, "Contacto: privacidad@gecontrol.mx. Domicilio: Carretera Sierra del Humo 211, colonia Bosques del Prado Sur, código postal 20130, Aguascalientes, Aguascalientes, México.")

heading(doc, "17", "CONSTANCIA DE PUESTA A DISPOSICIÓN")
para(doc, "Nombre de la persona titular: ______________________________________________")
para(doc, "Empresa / RFC relacionado: _________________________________________________")
para(doc, "Medio de puesta a disposición:  ☐ Plataforma   ☐ Correo   ☐ Impreso   ☐ Otro: __________")
para(doc, "Versión: 1.0    Fecha: __________________    Firma o aceptación: __________________________")
para(doc, "La firma acredita recepción o puesta a disposición del Aviso; no implica renuncia a derechos ni consentimiento para tratamientos que legalmente requieran manifestación separada.", italic=True)

doc.add_page_break()
heading(doc, "ANEXO INTERNO", "CONTROL DE VERSIONES PARA SUPERADMIN")
para(doc, "Esta página es de control interno y puede omitirse de la copia entregada.", italic=True)
grid(doc, ["Campo", "Fuente", "Regla"], [
    ("Responsable legal", "Perfil legal versionado.", "Debe resolverse al emitir."),
    ("Domicilio", "Domicilio para notificaciones.", "Completo y vigente."),
    ("Correo ARCO", "privacidad@gecontrol.mx", "Canal monitoreado."),
    ("Versión y vigencia", "Plantilla publicada.", "Inmutable al aceptar."),
    ("Proveedores", "Registro de encargados.", "Actualizar por nueva versión."),
    ("Plazos", "Política aprobada.", "No cambiar retroactivamente."),
    ("Aceptación", "Usuario, fecha, IP/request ID.", "Evidencia append-only."),
], [3000, 3400, 2960])
para(doc, "Superadmin deberá conservar la versión exacta puesta a disposición. Un cambio crea una nueva versión; no debe reemplazar silenciosamente la evidencia histórica.")

doc.save(DEST)
print(DEST)
