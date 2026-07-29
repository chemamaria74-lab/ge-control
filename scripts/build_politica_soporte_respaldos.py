from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Politica_Soporte_Respaldos_Conservacion_Exportacion_GE_Control.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.3, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    for kind, text in (("begin", None), (None, instruction), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText" if text == instruction else "w:t")
            node.text = text
            if text == instruction:
                node.set(qn("xml:space"), "preserve")
        run._r.append(node)


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=9.8)


def kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_width(cells[0], 2800)
        set_width(cells[1], 6560)
        shade(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.2, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.2)
    return table


def grid(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        set_width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8.8, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(value), size=font_size)
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.3)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.14
for style_name, size in (("Heading 1", 14), ("Heading 2", 11.5)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(11)
    style.paragraph_format.space_after = Pt(5)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  SOPORTE, RESPALDOS Y EXPORTACIÓN"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Política operativa v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("POLÍTICA DE SOPORTE, RESPALDOS,\nCONSERVACIÓN Y EXPORTACIÓN"), size=20.5, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(15)
set_font(subtitle.add_run("Portal Transporte · Anexo operativo del Contrato Marco SaaS"), size=12.5, color=MUTED)

kv_table(doc, [
    ("Versión", "{{politica_operativa_version}}"),
    ("Vigencia", "{{politica_operativa_fecha_vigencia}}"),
    ("Contrato Marco", "{{contrato_marco_folio}}"),
    ("Orden de Servicio", "{{orden_servicio_folio}}"),
    ("Cliente / RFC", "{{cliente_nombre_legal}} · {{suscripcion_rfc}}"),
    ("Plan", "{{plan_nombre}}"),
])

para(doc, "Esta Política establece los procedimientos operativos para solicitar soporte, conservar y recuperar información, obtener exportaciones y gestionar la terminación del Servicio. Forma parte del Contrato Marco y debe interpretarse junto con el SLA, la Orden de Servicio y el Anexo de Tratamiento y Seguridad.")

heading(doc, "1", "ALCANCE Y PRINCIPIOS")
para(doc, "Aplica a Portal Transporte y al Portal del Operador cuando esté contratado. Sus objetivos son dar trazabilidad a solicitudes, proteger la continuidad del Servicio, permitir al CLIENTE recuperar su información y evitar que los respaldos sean confundidos con un archivo histórico ilimitado.")
for item in [
    "El CLIENTE conserva la responsabilidad sobre la calidad, legalidad y oportunidad de sus datos.",
    "GE Control administra la infraestructura y asistencia dentro del alcance contratado.",
    "Los respaldos sirven para continuidad y recuperación general; no sustituyen la conservación propia del CLIENTE.",
    "Los plazos operativos pueden ampliarse por obligación legal, incidente, reclamación o defensa de derechos.",
]:
    bullet(doc, item)

heading(doc, "2", "CANALES Y HORARIO DE SOPORTE")
grid(doc, ["Concepto", "Condición"], [
    ("Recepción de solicitudes", "24 horas al día, 7 días a la semana, mediante los canales habilitados."),
    ("Atención humana ordinaria", "Lunes a viernes de 09:00 a 18:00, hora America/Mexico_City, excepto días de descanso obligatorio."),
    ("Canal principal", "{{correo_soporte}} o formulario/ticket dentro de la plataforma."),
    ("Incidente crítico", "Canal de emergencia indicado en {{canal_emergencia}}, disponible sólo para severidad P1."),
    ("Idioma", "Español."),
], [2800, 6560])
para(doc, "La recepción automática no significa resolución inmediata. Los tiempos de primera respuesta y actualizaciones se rigen por el SLA contratado y comienzan a computarse conforme a sus reglas de horario hábil, salvo cobertura especial expresamente contratada.")

heading(doc, "3", "CONTENIDO MÍNIMO DEL TICKET", page_break_before=True)
for item in [
    "Cliente, RFC afectado, usuario y medio de contacto.",
    "Descripción clara, fecha y hora aproximada, módulo y operación intentada.",
    "Folio, UUID, viaje o identificador relacionado, sin incluir contraseñas ni llaves privadas.",
    "Capturas, mensaje de error y pasos para reproducir, ocultando datos personales innecesarios.",
    "Impacto operativo, número estimado de usuarios y existencia de alternativa temporal.",
]:
    bullet(doc, item)
para(doc, "GE Control podrá solicitar información adicional. El tiempo durante el cual el ticket permanezca pendiente de datos, acceso o validación del CLIENTE no se imputará a los objetivos de respuesta o solución.")

heading(doc, "4", "CLASIFICACIÓN Y ATENCIÓN")
grid(doc, ["Prioridad", "Criterio", "Ejemplos"], [
    ("P1 Crítica", "Servicio general inaccesible o función esencial detenida sin alternativa.", "Acceso general caído; imposibilidad general atribuible a GE Control."),
    ("P2 Alta", "Función principal severamente afectada, con alcance relevante.", "Error repetido en un módulo principal; degradación importante."),
    ("P3 Media", "Afectación limitada o con alternativa razonable.", "Error individual, consulta o ajuste de configuración."),
    ("P4 Baja", "Solicitud informativa, mejora o cambio no urgente.", "Capacitación, explicación, propuesta de funcionalidad."),
], [1500, 3930, 3930], 8.1)
para(doc, "GE Control podrá reclasificar un ticket con explicación. Validaciones del SAT, rechazo de datos, indisponibilidad del PAC o un tercero y errores causados por configuración del CLIENTE no se consideran automáticamente una falla de GE Control.")

heading(doc, "5", "ALCANCE Y EXCLUSIONES DEL SOPORTE")
para(doc, "El soporte incluido comprende diagnóstico, orientación de uso, corrección de defectos reproducibles, atención de incidentes y ayuda razonable con funciones estándar.")
para(doc, "No incluye captura masiva, contabilidad, asesoría fiscal o jurídica, corrección de datos fuente, desarrollo personalizado, capacitación extensa, recuperación por errores del CLIENTE, integración no contratada, soporte a equipos o redes, ni trabajo sobre sistemas de terceros. Estos servicios podrán cotizarse por separado.")
para(doc, "El CLIENTE no deberá compartir contraseñas, certificados, llaves privadas o credenciales completas por ticket. Cuando un acceso temporal sea indispensable, se utilizará un mecanismo autorizado y deberá revocarse al concluir.")

heading(doc, "6", "MANTENIMIENTO Y COMUNICACIONES")
para(doc, "El mantenimiento programado se comunicará, cuando sea razonablemente posible, con al menos cuarenta y ocho horas de anticipación. La ventana ordinaria es el domingo de 00:00 a 04:00, hora America/Mexico_City.")
para(doc, "GE Control podrá ejecutar mantenimiento urgente sin aviso previo suficiente para contener una vulnerabilidad, evitar pérdida de datos o responder a una dependencia crítica. Informará el motivo y estado tan pronto como sea prudente.")

heading(doc, "7", "RESPALDOS Y RECUPERACIÓN", page_break_before=True)
para(doc, "GE Control mantendrá procedimientos de respaldo y recuperación adecuados a la arquitectura productiva. Como objetivo base, la pérdida máxima de datos recuperables (RPO) y el tiempo de recuperación (RTO) serán de hasta veinticuatro horas, sujetos a las exclusiones del SLA.")
grid(doc, ["Elemento", "Regla operativa"], [
    ("Finalidad", "Recuperación ante falla grave, corrupción, eliminación general o incidente de infraestructura."),
    ("Frecuencia objetivo", "Al menos diaria para la información productiva incluida en el esquema de respaldo."),
    ("Retención técnica", "{{dias_retencion_respaldos}} días en ciclos rotativos, salvo configuración del proveedor."),
    ("Protección", "Acceso restringido, separación lógica y protecciones de la infraestructura contratada."),
    ("Pruebas", "Revisión o prueba periódica proporcional al riesgo y después de cambios relevantes."),
    ("Restauración", "Priorizada para recuperar el Servicio completo o un conjunto coherente de datos."),
], [2500, 6860])
para(doc, "Los respaldos no son copias documentales navegables, no garantizan restaurar un registro individual y pueden sobrescribirse por rotación. La recuperación puntual solicitada por el CLIENTE estará sujeta a viabilidad, integridad, seguridad y cotización previa.")

heading(doc, "8", "CONSERVACIÓN Y BLOQUEO")
grid(doc, ["Categoría", "Plazo operativo base", "Regla especial"], [
    ("Documentos fiscales y trazabilidad", "Al menos 5 años.", "Se aplicará el plazo legal mayor."),
    ("Evidencias operativas", "24 meses.", "Puede ampliarse por incidente o reclamación."),
    ("Geolocalización por evento", "12 meses.", "No se amplía salvo incidencia, reclamación u obligación."),
    ("Prospectos no contratados", "12 meses.", "Desde la última interacción útil."),
    ("Registros de seguridad", "12 meses.", "Puede ampliarse durante una investigación."),
], [3100, 2200, 4060], 8.1)
para(doc, "Concluida la finalidad, la información se bloqueará cuando corresponda y posteriormente se eliminará o anonimizará. El CLIENTE deberá descargar oportunamente la información que necesite conservar por más tiempo.")

heading(doc, "9", "EXPORTACIONES DURANTE LA VIGENCIA")
para(doc, "El CLIENTE podrá utilizar las funciones de descarga disponibles. Según el módulo, la exportación podrá incluir XML, PDF, CSV, ZIP, JSON o archivos originales. El formato exacto depende de la naturaleza del dato y de la funcionalidad vigente.")
for item in [
    "Las exportaciones se limitan al RFC y permisos del usuario solicitante.",
    "No incluyen código fuente, secretos, modelos internos, datos de otros clientes ni una copia de la base de datos.",
    "La generación puede dividirse por periodo o tipo de información por seguridad y tamaño.",
    "El CLIENTE debe almacenar la descarga de forma segura y verificar su integridad y apertura.",
]:
    bullet(doc, item)
para(doc, "Una exportación asistida estándar se entregará razonablemente dentro de cinco días hábiles desde una solicitud completa. Volúmenes extraordinarios, reconstrucción histórica o formatos especiales requerirán estimación y, en su caso, cotización.")

doc.add_page_break()
heading(doc, "10", "SUSPENSIÓN POR FALTA DE PAGO")
para(doc, "Salvo condición distinta en la Orden de Servicio, se aplicará el siguiente esquema después del vencimiento de una factura no controvertida:")
grid(doc, ["Momento", "Medida"], [
    ("Días 1 a 9", "Avisos de cobro y continuidad ordinaria, sin renuncia al pago."),
    ("Desde el día 10", "Podrá bloquearse la creación de nuevas operaciones y el timbrado. Se procurará conservar consulta, exportación y cancelaciones fiscales necesarias, cuando sea técnica y legalmente posible."),
    ("Desde el día 30", "GE Control podrá terminar la suscripción y comenzar la ventana posterior de exportación."),
    ("Pago regularizado", "Restablecimiento razonable dentro de un día hábil desde que el pago sea identificado y validado."),
], [2200, 7160], 8.3)
para(doc, "La suspensión no elimina adeudos ni amplía los plazos fiscales del CLIENTE. GE Control podrá suspender inmediatamente accesos comprometidos, usos ilícitos o riesgos graves de seguridad, limitando la medida a lo razonablemente necesario.")

heading(doc, "11", "TERMINACIÓN Y VENTANA DE EXPORTACIÓN")
para(doc, "Al terminar una suscripción, el CLIENTE contará con treinta días naturales para solicitar o completar la exportación, salvo que la Orden de Servicio establezca un plazo mayor. Durante esa ventana el acceso podrá limitarse a lectura y descarga.")
para(doc, "Vencido el plazo, GE Control podrá deshabilitar el acceso e iniciar eliminación o anonimización de sistemas activos, respetando bloqueo, obligaciones legales, evidencia contractual, documentos fiscales y respaldos rotativos.")
para(doc, "La falta de descarga dentro de la ventana será responsabilidad del CLIENTE. Una recuperación posterior no se garantiza y, si fuera técnicamente posible, podrá generar cargos.")

heading(doc, "12", "SOLICITUD DE ELIMINACIÓN ANTICIPADA")
para(doc, "El CLIENTE podrá solicitar eliminación antes del plazo operativo cuando no exista deber de conservación, bloqueo, reclamación, investigación o necesidad contractual. La solicitud debe provenir de un contacto autorizado, identificar el RFC, alcance y confirmar que comprende su carácter irreversible.")
para(doc, "GE Control podrá requerir verificación reforzada y una segunda confirmación. La eliminación de sistemas activos no implica borrado inmediato de respaldos rotativos, los cuales quedarán aislados hasta expirar por ciclo.")

heading(doc, "13", "RESPONSABILIDADES DEL CLIENTE")
for item in [
    "Designar contactos vigentes de soporte, facturación, privacidad y seguridad.",
    "Mantener administradores y permisos correctos; revocar usuarios que ya no deban acceder.",
    "Conservar copias propias de información crítica y realizar exportaciones periódicas.",
    "Atender avisos, probar alternativas razonables y colaborar con el diagnóstico.",
    "Verificar obligaciones fiscales, laborales y documentales independientes del Servicio.",
]:
    bullet(doc, item)

heading(doc, "14", "CAMBIOS Y VERSIONADO")
para(doc, "GE Control podrá actualizar procedimientos sin disminuir materialmente los compromisos contratados. Los cambios materiales se comunicarán por plataforma, correo u otro medio registrado y se aplicarán prospectivamente.")
para(doc, "Cada Orden de Servicio conservará la versión aceptada. Un cambio de tarifa, plazo de exportación, horario o nivel de soporte deberá reflejarse en una nueva versión o documento modificatorio.")

heading(doc, "APÉNDICE A", "PARÁMETROS OPERATIVOS", page_break_before=True)
kv_table(doc, [
    ("Correo de soporte", "{{correo_soporte}}"),
    ("Canal P1", "{{canal_emergencia}}"),
    ("Horario humano", "Lunes a viernes, 09:00 a 18:00, America/Mexico_City"),
    ("Recepción", "24/7"),
    ("RPO / RTO objetivo", "Hasta 24 horas / hasta 24 horas"),
    ("Retención técnica de respaldos", "{{dias_retencion_respaldos}} días"),
    ("Ventana posterior de exportación", "30 días naturales"),
])

heading(doc, "APÉNDICE B", "ACEPTACIÓN")
para(doc, "Esta Política se acepta como parte de la Orden de Servicio. La firma electrónica, aceptación verificable o firma autógrafa producirá los efectos previstos en el Contrato Marco.")
grid(doc, ["POR EL CLIENTE", "POR GE CONTROL"], [
    ("Nombre: {{cliente_representante}}\nCargo: {{cliente_cargo}}\nFirma: __________________________\nFecha: {{cliente_fecha_firma}}",
     "María José Mejía Ornelas\nGE Control\nFirma: __________________________\nFecha: {{gecontrol_fecha_firma}}"),
], [4680, 4680], 9.0)

heading(doc, "APÉNDICE INTERNO", "CAMPOS PARA SUPERADMIN", page_break_before=True)
para(doc, "Este apéndice es operativo y deberá omitirse del PDF firmado. Las reglas legales permanecen versionadas; sólo se completan o seleccionan parámetros previamente aprobados.", italic=True)
grid(doc, ["Grupo", "Campos editables o seleccionables"], [
    ("Documento", "Versión, vigencia, Contrato Marco, Orden de Servicio, cliente, RFC y plan."),
    ("Soporte", "Correo, canal P1 y nivel contratado. El horario base sólo cambia mediante versión aprobada."),
    ("Respaldos", "Retención técnica compatible con el proveedor; RPO/RTO deben coincidir con el SLA."),
    ("Exportación", "Formato disponible, solicitud, estado, fecha de vencimiento, archivo, hash y persona que descargó."),
    ("Suspensión", "Factura, fecha de vencimiento, avisos, fecha de bloqueo, motivo, regularización y reactivación."),
    ("Terminación", "Fecha efectiva, inicio y fin de ventana, última exportación, bloqueo y orden de eliminación."),
], [2200, 7160], 8.3)

doc.core_properties.title = "Política de Soporte, Respaldos, Conservación y Exportación - GE Control"
doc.core_properties.subject = "Portal Transporte - Política operativa"
doc.core_properties.author = "GE Control"
doc.core_properties.keywords = "soporte, respaldos, conservación, exportación, SaaS, Portal Transporte"
doc.save(DEST)
print(DEST)
