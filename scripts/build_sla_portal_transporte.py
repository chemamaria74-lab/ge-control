from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Acuerdo_Niveles_Servicio_SLA_Portal_Transporte.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=10)


def kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        width(cells[0], 2800)
        width(cells[1], 6560)
        shade(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.5)
    return table


def grid(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=9, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(value), size=8.7)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  ACUERDO DE NIVELES DE SERVICIO"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("SLA Portal Transporte  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("ACUERDO DE NIVELES DE SERVICIO"), size=22, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
set_font(subtitle.add_run("Portal Transporte · Anexo del Contrato Marco SaaS"), size=13, color=MUTED)

kv_table(doc, [
    ("Versión", "{{sla_version}}"),
    ("Vigencia", "{{sla_fecha_vigencia}}"),
    ("Contrato Marco", "{{contrato_marco_folio}}"),
    ("Orden de Servicio", "{{orden_servicio_folio}}"),
    ("Cliente / RFC", "{{cliente_nombre_legal}} · {{suscripcion_rfc}}"),
    ("Nivel contratado", "{{sla_nivel_contratado}}"),
])

para(doc, "El presente Acuerdo de Niveles de Servicio (“SLA”) forma parte del Contrato Marco y establece objetivos medibles de disponibilidad, soporte, continuidad y atención de incidentes. Los términos con mayúscula inicial tendrán el significado del Contrato Marco.")

heading(doc, "1", "ALCANCE")
para(doc, "Este SLA aplica a las funciones productivas de Portal Transporte expresamente habilitadas en la Orden de Servicio, incluyendo acceso administrativo, APIs propias, catálogos, viajes, generación documental y consulta de información.")
para(doc, "El Portal del Operador queda cubierto cuando exista un addon activo. Los servicios de certificación del PAC, conectividad pública, dispositivos, navegadores, correo y plataformas de terceros se consideran dependencias y se sujetan a las exclusiones previstas.")

heading(doc, "2", "DISPONIBILIDAD MENSUAL")
para(doc, "GE Control establece como objetivo una Disponibilidad Mensual de 99.5% para el Servicio principal.")
grid(doc, ["Concepto", "Definición"], [
    ("Periodo de medición", "Mes calendario, zona America/Mexico_City."),
    ("Servicio disponible", "Usuarios autorizados pueden autenticarse y ejecutar funciones principales sin error general atribuible a GE Control."),
    ("Minutos del periodo", "Minutos totales del mes menos exclusiones válidas."),
    ("Indisponibilidad", "Minutos completos de afectación general confirmada y atribuible a GE Control."),
    ("Fórmula", "((Minutos aplicables − indisponibilidad) / minutos aplicables) × 100."),
], [2600, 6760])

para(doc, "Una degradación parcial se contabilizará cuando impida materialmente una función principal a una proporción significativa de usuarios. Errores individuales de datos, validaciones fiscales o configuración no constituyen por sí solos indisponibilidad.")

heading(doc, "3", "EXCLUSIONES DE DISPONIBILIDAD")
for item in [
    "Mantenimiento programado comunicado con al menos 48 horas de anticipación, cuando sea razonablemente posible.",
    "Mantenimiento urgente necesario para contener una vulnerabilidad o evitar daño.",
    "Fallas de SW Sapien, Supabase, Render, DNS, correo, telecomunicaciones u otro tercero fuera del control razonable de GE Control, siempre que GE Control gestione y escale el incidente.",
    "Fallas de Internet, red, dispositivo, navegador, antivirus, firewall o configuración del CLIENTE.",
    "Uso contrario a documentación, límites, contrato o instrucciones de soporte.",
    "Datos fiscales incorrectos, certificados vencidos, saldos PAC, rechazos SAT/PAC o validaciones propias del documento.",
    "Suspensión por falta de pago, orden de autoridad, riesgo de seguridad o actuación solicitada por el CLIENTE.",
    "Fuerza mayor, ataques generalizados o eventos imprevisibles que no hubieran podido evitarse con medidas razonables.",
    "Entornos de prueba, beta, demostración, funciones experimentales o integraciones no incluidas en la Orden de Servicio.",
]:
    bullet(doc, item)

heading(doc, "4", "MANTENIMIENTO")
para(doc, "La ventana preferente de mantenimiento programado será domingo de 00:00 a 04:00 horas, zona America/Mexico_City. GE Control podrá utilizar otra ventana cuando la naturaleza del cambio lo requiera.")
para(doc, "El aviso indicará alcance, inicio estimado, duración y afectación prevista. Los mantenimientos urgentes podrán ejecutarse sin el plazo ordinario; GE Control informará tan pronto como sea razonablemente posible.")

heading(doc, "5", "CANALES Y HORARIO DE SOPORTE")
kv_table(doc, [
    ("Recepción de solicitudes", "24 horas, todos los días, mediante canales habilitados"),
    ("Atención humana ordinaria", "Lunes a viernes, 09:00 a 18:00, America/Mexico_City, excepto días inhábiles oficiales"),
    ("Correo general", "{{soporte_correo}}"),
    ("Canal crítico", "{{soporte_canal_critico}}"),
    ("Idioma", "Español"),
])
para(doc, "La recepción 24/7 significa que el sistema puede recibir, foliar y acusar una solicitud en cualquier momento. No implica atención humana o resolución inmediata fuera del horario ordinario, salvo que la Orden de Servicio contrate una guardia especial.")

heading(doc, "6", "PRIORIDADES Y PRIMERA RESPUESTA")
grid(doc, ["Prioridad", "Criterio", "Primera respuesta objetivo", "Actualización"], [
    ("P1 Crítica", "Servicio general inaccesible, pérdida activa de datos o riesgo grave de seguridad sin alternativa.", "Hasta 4 horas hábiles", "Cada 4 horas hábiles"),
    ("P2 Alta", "Función principal afectada para varios usuarios; alternativa limitada.", "Hasta 8 horas hábiles", "Cada día hábil"),
    ("P3 Media", "Afectación individual o parcial con alternativa disponible.", "Hasta 2 días hábiles", "Al existir avance material"),
    ("P4 Baja", "Consulta, configuración, mejora o defecto cosmético.", "Hasta 3 días hábiles", "Según planificación"),
], [1400, 3900, 2200, 1860])

para(doc, "“Primera respuesta” significa confirmación humana, clasificación y solicitud de información o acción inicial. No es una garantía de solución. Los plazos se pausan mientras GE Control espera datos, acceso, reproducción o autorización del CLIENTE.")

heading(doc, "7", "GESTIÓN DE INCIDENTES")
para(doc, "GE Control registrará los incidentes, asignará prioridad, investigará causa, aplicará mitigación y comunicará avances razonables. La prioridad puede ajustarse conforme a evidencia y alcance real.")
para(doc, "Para P1 atribuibles a GE Control se preparará, cuando resulte material, un resumen posterior con impacto, duración, causa conocida, mitigación y acciones preventivas. El informe podrá omitir información que comprometa seguridad o confidencialidad.")

heading(doc, "8", "OBLIGACIONES DEL CLIENTE")
for item in [
    "Designar contactos autorizados y mantener actualizados sus medios.",
    "Reportar con fecha, usuario, RFC, viaje, captura, mensaje y pasos de reproducción, evitando compartir contraseñas o llaves.",
    "Colaborar en pruebas y confirmar la recuperación.",
    "Mantener dispositivos, navegadores, conectividad, certificados y datos fiscales en condiciones adecuadas.",
    "No duplicar tickets por el mismo evento ni clasificar como crítico un caso que no cumple el criterio.",
    "Descargar y conservar los documentos que requiera su política interna.",
]:
    bullet(doc, item)

heading(doc, "9", "CONTINUIDAD, RESPALDOS Y RECUPERACIÓN")
para(doc, "GE Control mantendrá mecanismos razonables de respaldo y recuperación conforme a su arquitectura y proveedores. Como objetivos operativos iniciales, no sujetos a créditos:")
grid(doc, ["Objetivo", "Valor base", "Alcance"], [
    ("RPO", "Hasta 24 horas", "Pérdida máxima objetivo ante desastre catastrófico."),
    ("RTO", "Hasta 24 horas", "Objetivo para restablecer funciones esenciales tras declarar desastre."),
    ("Prueba de restauración", "Periódica", "Conforme a la política interna y capacidad contratada."),
], [1800, 2200, 5360])
para(doc, "RPO y RTO son objetivos de continuidad, no garantías de recuperación de cada registro ni sustituyen las obligaciones de conservación del CLIENTE. Un incidente ordinario no activa automáticamente el procedimiento de desastre.")

heading(doc, "10", "CRÉDITOS DE SERVICIO")
para(doc, "Si la Disponibilidad Mensual atribuible a GE Control es inferior al objetivo, el CLIENTE podrá solicitar un crédito sobre la cuota recurrente mensual de la suscripción afectada:")
grid(doc, ["Disponibilidad mensual", "Crédito"], [
    ("99.00% a 99.49%", "5%"),
    ("98.00% a 98.99%", "10%"),
    ("Menor a 98.00%", "20%"),
], [5600, 3760])
para(doc, "El crédito máximo por mes es 20% de la cuota recurrente del Servicio afectado; no incluye IVA, implementación, PAC, addon no afectado ni servicios profesionales. Se aplicará a una factura futura y no se paga en efectivo.")

heading(doc, "11", "SOLICITUD Y VALIDACIÓN DE CRÉDITOS")
para(doc, "El CLIENTE deberá solicitar el crédito dentro de los diez días naturales siguientes al cierre del mes, indicando suscripción, fechas, duración y tickets relacionados. GE Control contrastará monitoreo, registros, exclusiones y afectación.")
para(doc, "El CLIENTE debe encontrarse al corriente. Los créditos son el remedio económico exclusivo por incumplimiento del objetivo de disponibilidad, sin limitar derechos que legalmente no puedan excluirse ni responsabilidades por dolo, culpa grave o supuestos previstos en el Contrato Marco.")

heading(doc, "12", "DEPENDENCIAS FISCALES Y TERCEROS")
para(doc, "GE Control gestiona integraciones, pero no controla los tiempos ni decisiones del SAT, PAC o proveedores. Un documento rechazado por datos, reglas, certificados o disponibilidad externa no constituye incumplimiento automático del SLA.")
para(doc, "GE Control procurará identificar si el origen es propio, del CLIENTE o de un tercero; escalará al proveedor cuando corresponda y comunicará alternativas disponibles, sin asumir obligaciones superiores a las contratadas con dicho tercero.")

heading(doc, "13", "SEGURIDAD Y NOTIFICACIONES")
para(doc, "Los eventos de seguridad se atenderán según riesgo. Cuando una vulneración afecte significativamente derechos o datos, GE Control notificará conforme a la ley y al Anexo de Tratamiento y Seguridad. La investigación forense, contención y preservación de evidencia pueden limitar temporalmente el detalle comunicado.")

heading(doc, "14", "VIGENCIA, CAMBIOS Y PRELACIÓN")
para(doc, "Este SLA inicia en {{sla_fecha_vigencia}} y permanece vigente mientras la Orden de Servicio correspondiente esté activa. Los cambios materiales requieren nueva versión y no modificarán retroactivamente periodos cerrados.")
para(doc, "En conflicto, el Contrato Marco rige responsabilidad general; la Orden de Servicio rige el nivel contratado; y este SLA rige métricas y atención. Una excepción deberá constar por escrito.")

heading(doc, "15", "ACEPTACIÓN")
para(doc, "Las PARTES aceptan el presente SLA como anexo del Contrato Marco y Orden de Servicio indicados.")
sig = doc.add_table(rows=4, cols=2)
sig.style = "Table Grid"
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
sig.autofit = False
rows = [
    ("GE CONTROL", "EL CLIENTE"),
    ("{{proveedor_nombre_legal}}", "{{cliente_nombre_legal}}"),
    ("{{proveedor_representante}}", "{{cliente_representante}}"),
    ("Firma/aceptación: __________________", "Firma/aceptación: __________________"),
]
for row, values in zip(sig.rows, rows):
    for idx, value in enumerate(values):
        width(row.cells[idx], 4680)
        p = row.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(7)
        set_font(p.add_run(value), size=9.5, bold=row is sig.rows[0])

doc.add_page_break()
heading(doc, "ANEXO INTERNO", "VARIABLES PARA SUPERADMIN")
para(doc, "Esta página es de control interno y puede omitirse del PDF firmado.", italic=True)
grid(doc, ["Variable", "Uso", "Control"], [
    ("{{sla_version}}", "Versión jurídica publicada.", "Inmutable al emitir"),
    ("{{sla_nivel_contratado}}", "Nivel base o especial.", "Tomado de Orden de Servicio"),
    ("{{soporte_correo}}", "Canal ordinario.", "Perfil legal/operativo"),
    ("{{soporte_canal_critico}}", "Canal de incidentes P1.", "Restringido a contactos"),
    ("Disponibilidad 99.5%", "Objetivo mensual.", "Cláusula protegida"),
    ("Créditos 5/10/20%", "Remedio por disponibilidad.", "Cláusula protegida"),
    ("RPO/RTO 24h", "Objetivos de continuidad.", "No sujetos a crédito"),
], [3000, 3600, 2760])
para(doc, "Una versión enviada o aceptada debe conservar valores resueltos, folio, hash, fecha, destinatarios y evidencia. Los porcentajes, tiempos y exclusiones no deben editarse cliente por cliente sin aprobar una versión especial.")

doc.save(DEST)
print(DEST)
