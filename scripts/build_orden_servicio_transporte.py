from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Orden_de_Servicio_Portal_Transporte_Plantilla.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def p(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.12
    set_font(para.add_run(text), italic=italic)
    return para


def heading(doc, text, level=1):
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.keep_with_next = True
    set_font(para.add_run(text), size=14 if level == 1 else 12, bold=True, color=WINE)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(.42)
    para.paragraph_format.first_line_indent = Inches(-.2)
    para.paragraph_format.space_after = Pt(3)
    set_font(para.add_run(text), size=10)


def kv_table(doc, rows, widths=(2800, 6560)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        width(cells[0], widths[0])
        width(cells[1], widths[1])
        shade(cells[0], LIGHT)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.5, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.5)
    return table


def grid(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(para.add_run(label), size=9, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cells[idx].paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            set_font(para.add_run(value), size=8.8)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  ORDEN DE SERVICIO"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Orden de Servicio versionada  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(22)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("ORDEN DE SERVICIO"), size=24, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
set_font(subtitle.add_run("Portal Transporte · Suscripción por RFC"), size=13, color=MUTED)

kv_table(doc, [
    ("Folio", "{{orden_servicio_folio}}"),
    ("Versión", "{{orden_servicio_version}}"),
    ("Contrato Marco", "{{contrato_marco_folio}} · versión {{contrato_marco_version}}"),
    ("Cotización aceptada", "{{cotizacion_folio}} · versión {{cotizacion_version}}"),
    ("Fecha de emisión", "{{fecha_emision}}"),
    ("Estatus", "{{estado_documento}}"),
])

heading(doc, "1. PARTES Y SUSCRIPCIÓN")
kv_table(doc, [
    ("Proveedor", "{{proveedor_nombre_legal}} · RFC {{proveedor_rfc}} · GE Control"),
    ("Cliente contractual", "{{cliente_nombre_legal}} · RFC {{cliente_rfc}}"),
    ("RFC de la suscripción", "{{suscripcion_rfc}}"),
    ("Razón social operativa", "{{suscripcion_razon_social}}"),
    ("Tenant", "{{tenant_reference}}"),
    ("Suscripción", "{{subscription_reference}}"),
    ("Contacto contractual", "{{cliente_contacto_nombre}} · {{cliente_contacto_correo}}"),
])

p(doc, "Esta Orden de Servicio se incorpora al Contrato Marco identificado y regula exclusivamente la suscripción correspondiente al RFC señalado. Otro RFC requiere su propia Orden de Servicio.")

heading(doc, "2. PLAN CONTRATADO")
grid(doc, ["Concepto", "Valor contratado", "Regla"], [
    ("Plan", "{{plan_nombre}} · versión {{plan_version}}", "Snapshot al aceptar"),
    ("Vehículos activos", "{{limite_vehiculos_activos}}", "Simultáneos; no incluye remolques"),
    ("Viajes fiscales", "{{limite_viajes_mensuales}} por mes", "Mes calendario · America/Mexico_City"),
    ("Administradores", "{{limite_administradores}}", "Individuales; invitaciones pendientes cuentan"),
    ("Operadores PIN", "Ilimitados", "Aislados por RFC"),
    ("Portal del Operador", "{{portal_operador_estado}}", "Addon independiente"),
], [2700, 3300, 3360])

heading(doc, "3. FUNCIONES INCLUIDAS")
p(doc, "La suscripción principal incluye, dentro de sus límites y configuración:")
for item in [
    "Viajes, clientes, rutas, orígenes, destinos, mercancías y catálogos operativos.",
    "Choferes, vehículos, remolques y expedientes relacionados.",
    "Generación, validación, timbrado, consulta y descarga de Carta Porte.",
    "Generación del CFDI de ingreso únicamente después de existir Carta Porte previa.",
    "XML, PDF, liquidaciones, evidencias, documentos y reportes disponibles.",
    "Actualizaciones generales, respaldo y soporte básico conforme a los anexos.",
]:
    bullet(doc, item)
p(doc, "No se venden estas funciones como módulos pequeños independientes. El Portal del Operador sí se contrata o concede como complemento separado.")

heading(doc, "4. PORTAL DEL OPERADOR")
kv_table(doc, [
    ("Modalidad", "{{portal_operador_modalidad}}"),
    ("Precio", "{{portal_operador_precio}} {{moneda}} más IVA"),
    ("Inicio", "{{portal_operador_inicio}}"),
    ("Fin de prueba/promoción", "{{portal_operador_promocion_fin}}"),
    ("Motivo de inclusión", "{{portal_operador_included_reason}}"),
    ("Estado al vencimiento", "{{portal_operador_expiration_action}}"),
])
p(doc, "La prueba o promoción no se renueva automáticamente salvo autorización. El permiso debe validarse en backend y puede suspenderse al vencer, sin eliminar información histórica.")

heading(doc, "5. PRECIO Y DESCUENTOS")
grid(doc, ["Partida", "Precio de lista", "Descuento", "Importe acordado"], [
    ("Suscripción {{plan_nombre}}", "{{plan_precio_lista}}", "{{plan_descuento}}", "{{plan_precio_acordado}}"),
    ("Portal del Operador", "{{portal_precio_lista}}", "{{portal_descuento}}", "{{portal_precio_acordado}}"),
    ("Implementación", "{{implementacion_lista}}", "{{implementacion_descuento}}", "{{implementacion_acordada}}"),
    ("Otros", "{{otros_lista}}", "{{otros_descuento}}", "{{otros_acordado}}"),
], [2700, 2200, 2100, 2360])

kv_table(doc, [
    ("Subtotal", "{{subtotal}} {{moneda}}"),
    ("IVA", "{{iva}} {{moneda}}"),
    ("Total inicial", "{{total_inicial}} {{moneda}}"),
    ("Total recurrente", "{{total_recurrente}} {{moneda}} más IVA"),
    ("Periodicidad", "{{periodicidad}}"),
    ("Forma y plazo de pago", "{{condiciones_pago}}"),
])

p(doc, "Todo descuento debe conservar tipo, base, motivo, aprobador y vigencia. Ningún descuento se presume permanente. El importe acordado de esta versión no cambia si posteriormente se modifica el catálogo.")

heading(doc, "6. VIGENCIA Y ACTIVACIÓN")
kv_table(doc, [
    ("Fecha de aceptación", "{{fecha_aceptacion}}"),
    ("Fecha de activación", "{{fecha_activacion}}"),
    ("Inicio del periodo facturable", "{{inicio_periodo_facturable}}"),
    ("Fin del periodo inicial", "{{fin_periodo_inicial}}"),
    ("Renovación", "{{regla_renovacion}}"),
    ("Aviso de no renovación", "Al menos 10 días naturales antes del siguiente periodo"),
])
p(doc, "La aceptación contractual no activa por sí sola el servicio. El periodo facturable comienza en la fecha de activación, después de pago o autorización expresa de GE Control.")

heading(doc, "7. VIAJES FISCALES Y LÍMITES")
for item in [
    "Una Carta Porte con UUID válido consume exactamente un viaje fiscal.",
    "El CFDI de ingreso relacionado no consume un segundo viaje.",
    "Intentos sin UUID y errores del PAC sin UUID no consumen.",
    "La cancelación no devuelve el viaje.",
    "Una nueva Carta Porte de sustitución con nuevo UUID consume otro viaje, salvo ajuste compensatorio autorizado por error interno comprobado.",
    "80% genera alerta; 90% genera alerta urgente; 100% bloquea nuevos timbrados de Carta Porte.",
    "No se bloquean consulta, descarga, exportación ni cancelaciones fiscales.",
    "No existen paquetes adicionales: el aumento ordinario de capacidad requiere cambio de plan.",
]:
    bullet(doc, item)
p(doc, "Un override temporal sólo puede concederse por Superadmin, debe indicar motivo, capacidad, responsable y vencimiento, y no altera permanentemente el plan.")

heading(doc, "8. CONDICIONES PARTICULARES")
kv_table(doc, [
    ("Plan legado", "{{legacy_flag}}"),
    ("Grandfathered", "{{grandfathered_flag}}"),
    ("Condición especial", "{{special_terms}}"),
    ("Override inicial", "{{initial_override}}"),
    ("Notas aprobadas", "{{approved_notes}}"),
])
p(doc, "Las notas comerciales no modifican cláusulas protegidas del Contrato Marco. Toda excepción debe aparecer expresamente en esta sección y contar con aprobación registrada.")

heading(doc, "9. DOCUMENTOS INCORPORADOS")
for item in [
    "Contrato Marco SaaS y su versión indicada.",
    "Acuerdo de Niveles de Servicio.",
    "Aviso de Privacidad aplicable.",
    "Anexo de Tratamiento y Seguridad de Datos.",
    "Política de Soporte, Respaldos, Conservación y Exportación.",
    "Reglas de Viajes Fiscales, Timbrado y Cancelaciones.",
    "Anexo de Geolocalización cuando se habilite el Portal del Operador.",
]:
    bullet(doc, item)

heading(doc, "10. ACEPTACIÓN")
p(doc, "Las PARTES aceptan esta Orden de Servicio y reconocen que sus datos comerciales constituyen un snapshot inmutable de la cotización aceptada.")
sig = doc.add_table(rows=4, cols=2)
sig.style = "Table Grid"
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
sig.autofit = False
signature_rows = [
    ("GE CONTROL", "EL CLIENTE"),
    ("{{proveedor_nombre_legal}}", "{{cliente_nombre_legal}}"),
    ("{{proveedor_representante}}", "{{cliente_representante}}"),
    ("Firma/aceptación: __________________", "Firma/aceptación: __________________"),
]
for row, values in zip(sig.rows, signature_rows):
    for idx, value in enumerate(values):
        width(row.cells[idx], 4680)
        para = row.cells[idx].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(7)
        para.paragraph_format.space_after = Pt(7)
        set_font(para.add_run(value), size=9.5, bold=row is sig.rows[0])

doc.add_page_break()
heading(doc, "ANEXO INTERNO · VARIABLES DE SUPERADMIN")
p(doc, "Esta página es de control interno y puede omitirse del PDF firmado.", italic=True)
grid(doc, ["Grupo", "Variables", "Regla de control"], [
    ("Documento", "{{orden_servicio_*}}, {{contrato_marco_*}}, {{cotizacion_*}}", "Versión emitida inmutable"),
    ("Suscripción", "{{subscription_*}}, {{suscripcion_rfc}}, {{plan_*}}", "Una principal activa por RFC/producto"),
    ("Capacidad", "{{limite_vehiculos_*}}, {{limite_viajes_*}}, {{limite_admin_*}}", "Snapshot + entitlements"),
    ("Precios", "{{*_precio_lista}}, {{*_descuento}}, {{*_acordado}}", "Conservar base y aprobador"),
    ("Portal", "{{portal_operador_*}}", "Trial/promoción con vencimiento"),
    ("Vigencia", "{{fecha_*}}, {{periodicidad}}, {{regla_renovacion}}", "Zona America/Mexico_City"),
], [2300, 4300, 2760])

heading(doc, "Reglas de emisión", level=2)
for item in [
    "Un borrador puede editarse; una versión enviada o aceptada no.",
    "Cambiar una versión enviada crea version_number + 1.",
    "El PDF debe contener valores resueltos, folio, versión, fecha y hash.",
    "Guardar actor, destinatario, evidencia de envío, aceptación e IP/request ID cuando aplique.",
    "No regenerar silenciosamente una Orden de Servicio firmada.",
]:
    bullet(doc, item)

doc.save(DEST)
print(DEST)
