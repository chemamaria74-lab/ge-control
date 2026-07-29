from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Anexo_Especial_Geolocalizacion_Operador_GE_Control.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.3, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=9.8)


def kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_width(cells[0], 2800)
        set_width(cells[1], 6560)
        shade(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.2, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.2)
    return table


def grid(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        set_width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8.8, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(value), size=font_size)
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.3)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.14
for style_name, size in (("Heading 1", 14), ("Heading 2", 11.5)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(11)
    style.paragraph_format.space_after = Pt(5)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  GEOLOCALIZACIÓN DEL OPERADOR"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Anexo de geolocalización v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("ANEXO ESPECIAL DE GEOLOCALIZACIÓN\nDEL OPERADOR"), size=21, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(15)
set_font(subtitle.add_run("Portal del Operador · Anexo del Contrato Marco SaaS"), size=12.5, color=MUTED)

kv_table(doc, [
    ("Versión", "{{geo_version}}"),
    ("Vigencia", "{{geo_fecha_vigencia}}"),
    ("Contrato Marco", "{{contrato_marco_folio}}"),
    ("Orden de Servicio", "{{orden_servicio_folio}}"),
    ("Cliente / Responsable", "{{cliente_nombre_legal}}"),
    ("RFC de la suscripción", "{{suscripcion_rfc}}"),
    ("Contacto de privacidad", "{{cliente_contacto_privacidad}}"),
])

para(doc, "Este Anexo regula la captura de ubicación precisa asociada a eventos operativos del viaje mediante el Portal del Operador. Su finalidad es impedir el rastreo excesivo, documentar las obligaciones del CLIENTE y delimitar el tratamiento que GE Control realiza como proveedor tecnológico.")

heading(doc, "1", "ROLES Y RESPONSABILIDAD")
para(doc, "El CLIENTE determina las finalidades laborales y operativas de la geolocalización y actúa como Responsable de los datos de sus operadores. GE Control actúa como Encargado cuando recibe, aloja y muestra la ubicación por instrucciones del CLIENTE.")
para(doc, "El CLIENTE es el único patrón, contratante o coordinador del operador. GE Control no dirige personal, asigna sanciones, evalúa desempeño laboral ni decide consecuencias disciplinarias.")

heading(doc, "2", "ALCANCE TÉCNICO APROBADO")
para(doc, "La configuración ordinaria captura una ubicación puntual únicamente cuando el operador registra voluntariamente una acción concreta en el Portal y el dispositivo concede permiso de ubicación.")
grid(doc, ["Característica", "Configuración comprometida"], [
    ("Modalidad", "Captura puntual asociada a un evento; no seguimiento continuo."),
    ("Activación", "Acción visible del operador y solicitud de permiso del sistema operativo."),
    ("Segundo plano", "Deshabilitado en la configuración ordinaria."),
    ("Fuera del viaje", "No se solicita ni captura ubicación para fines del Portal."),
    ("Frecuencia", "Una evidencia por evento activado; reintentos sólo cuando el operador lo solicita."),
    ("Precisión", "La disponible en el dispositivo; puede variar por señal, equipo y entorno."),
], [2800, 6560])
para(doc, "Cualquier función futura de rastreo continuo, segundo plano, geocercas automáticas o telemetría recurrente requerirá evaluación jurídica y de impacto, una nueva versión de este Anexo, aviso específico y configuración separada. No queda autorizada por este documento.")

heading(doc, "3", "EVENTOS Y FINALIDADES PERMITIDAS")
grid(doc, ["Evento configurable", "Finalidad permitida"], [
    ("Inicio de viaje", "Acreditar el punto y momento declarados de inicio."),
    ("Llegada a origen", "Documentar arribo para carga o recolección."),
    ("Carga concluida", "Respaldar el evento operativo informado."),
    ("Incidencia", "Ubicar el sitio declarado para coordinación y evidencia."),
    ("Llegada a destino", "Documentar arribo para descarga o entrega."),
    ("Entrega concluida", "Respaldar cierre del evento o viaje."),
], [3300, 6060])
para(doc, "El CLIENTE podrá habilitar sólo los eventos necesarios. No deberá reutilizar la ubicación para publicidad, vigilancia personal, perfiles ajenos al viaje, investigación de actividades privadas ni finalidades incompatibles.")

heading(doc, "4", "DATOS CAPTURADOS")
for item in [
    "latitud y longitud;",
    "precisión reportada por el dispositivo;",
    "fecha y hora del evento;",
    "tipo de evento, viaje, RFC y operador asociado;",
    "identificador técnico, usuario, IP o registro de seguridad necesario; y",
    "observaciones o evidencia que el operador decida adjuntar dentro del flujo.",
]:
    bullet(doc, item)
para(doc, "Portal Transporte no necesita acceder a contactos, llamadas, mensajes, fotografías privadas, micrófono ni historial general de ubicaciones del dispositivo. El CLIENTE no deberá exigir permisos ajenos a las funciones habilitadas.")

heading(doc, "5", "INFORMACIÓN AL OPERADOR Y BASE JURÍDICA")
para(doc, "Antes de habilitar la función, el CLIENTE deberá entregar al operador un aviso de privacidad que identifique al Responsable, datos tratados, finalidades, medios de contacto, conservación, destinatarios y mecanismo para ejercer derechos.")
para(doc, "El CLIENTE deberá determinar y documentar la base jurídica aplicable considerando la relación laboral o contractual, necesidad, proporcionalidad y alternativas menos invasivas. El consentimiento no deberá utilizarse como fórmula automática cuando exista subordinación o cuando no pueda otorgarse libremente.")
para(doc, "El permiso concedido en iOS, Android o navegador es una autorización técnica del dispositivo. Por sí solo no acredita que el CLIENTE haya cumplido sus obligaciones de información, licitud o proporcionalidad.")

heading(doc, "6", "OBLIGACIONES LABORALES DEL CLIENTE")
for item in [
    "Informar por escrito la política de uso, eventos, horarios, responsables y consecuencias operativas.",
    "Limitar la captura al viaje y jornada o servicio efectivamente asignados.",
    "Proporcionar una alternativa razonable cuando falle el GPS, el permiso o el dispositivo.",
    "No sancionar automáticamente al operador por una lectura aislada, inexacta o ausente.",
    "Permitir aclaraciones y valorar otras evidencias antes de una decisión adversa.",
    "Regular el uso de dispositivos personales, conectividad, costos y soporte conforme a su relación laboral.",
    "Capacitar a supervisores y restringir la consulta a personal con necesidad operativa.",
]:
    bullet(doc, item)
para(doc, "El CLIENTE deberá revisar con su abogado laboral la política interna, contratos, reglamento y avisos aplicables. Este Anexo no autoriza vigilancia fuera de jornada ni modifica derechos del operador.")

heading(doc, "7", "PROHIBICIONES DE USO")
grid(doc, ["Prohibido", "Alcance"], [
    ("Rastreo oculto o continuo", "No activar capturas recurrentes o en segundo plano sin un nuevo marco jurídico y técnico."),
    ("Ubicación fuera del servicio", "No solicitar eventos simulados para conocer domicilio, descanso o actividades privadas."),
    ("Decisión automatizada exclusiva", "No sancionar, despedir, reducir pago o calificar desempeño únicamente por coordenadas."),
    ("Acceso indiscriminado", "No habilitar ubicación a usuarios sin función y RFC autorizados."),
    ("Difusión pública", "No publicar mapas, recorridos o ubicaciones identificables."),
    ("Finalidad discriminatoria", "No inferir salud, religión, afiliación, vida privada u otras categorías sensibles."),
], [3000, 6360], 8.2)

heading(doc, "8", "EXACTITUD Y VALOR DE LA EVIDENCIA")
para(doc, "La ubicación puede ser inexacta, tardía o inexistente por edificios, clima, cobertura, configuración, batería, permisos, reloj, GPS, navegador o intervención del usuario. Una coordenada acredita lo reportado por el dispositivo, no necesariamente la conducta, intención o permanencia de una persona.")
para(doc, "La geolocalización deberá evaluarse junto con sellos temporales, documentos, comunicaciones y demás evidencia. El operador podrá reportar una inconsistencia y el CLIENTE deberá conservar la aclaración vinculada al evento.")

heading(doc, "9", "ACCESOS Y TRAZABILIDAD")
para(doc, "Sólo usuarios autorizados del CLIENTE, limitados por cliente y RFC, podrán consultar ubicación. GE Control podrá acceder cuando sea indispensable para soporte, seguridad, incidente o requerimiento legal.")
para(doc, "La plataforma conservará trazabilidad razonable de creación, consulta administrativa, exportación, corrección técnica y eliminación, incluyendo actor, fecha, finalidad o ticket cuando corresponda.")

heading(doc, "10", "CONSERVACIÓN Y ELIMINACIÓN")
para(doc, "La ubicación asociada a eventos se conservará durante doce meses desde su captura. Podrá mantenerse por más tiempo únicamente cuando esté vinculada a una incidencia, reclamación, investigación, obligación legal o defensa de derechos.")
para(doc, "Cumplido el plazo, se eliminará o anonimizará en sistemas activos conforme a los procesos técnicos. Las copias en respaldos rotativos expiran por ciclo y permanecen aisladas de uso ordinario.")
para(doc, "El CLIENTE deberá evitar exportaciones indiscriminadas y eliminar sus copias cuando concluya la finalidad o el plazo que haya informado.")

heading(doc, "11", "DERECHOS DEL OPERADOR")
para(doc, "El operador podrá ejercer acceso, rectificación, cancelación u oposición ante el CLIENTE mediante {{cliente_contacto_privacidad}}. También podrá solicitar aclaración de un evento, conocer quién puede consultarlo y limitar usos incompatibles.")
para(doc, "Si GE Control recibe una solicitud relacionada con Datos del Cliente, la remitirá al contacto registrado y prestará asistencia razonable. GE Control atenderá directamente los tratamientos en los que actúe como Responsable a través de privacidad@gecontrol.mx.")

heading(doc, "12", "SEGURIDAD E INCIDENTES")
for item in [
    "autenticación y permisos por rol;",
    "segregación lógica por cliente y RFC;",
    "transmisión protegida mediante TLS en servicios compatibles;",
    "almacenamiento privado y acceso administrativo restringido;",
    "registros y procedimientos de atención de incidentes; y",
    "respaldos y eliminación conforme a las políticas contractuales.",
]:
    bullet(doc, item)
para(doc, "GE Control notificará al CLIENTE los incidentes confirmados conforme al Anexo de Tratamiento y Seguridad. El CLIENTE decidirá y realizará las comunicaciones a operadores o autoridades como Responsable.")

heading(doc, "13", "DISPOSITIVOS Y CONTINUIDAD")
para(doc, "El CLIENTE definirá si proporciona equipo corporativo o autoriza un dispositivo personal. Deberá documentar requisitos mínimos, conectividad, actualizaciones, soporte y separación razonable entre información laboral y privada.")
para(doc, "Si no puede obtenerse ubicación, el Portal deberá permitir registrar o reportar la incidencia por un mecanismo alternativo cuando esté disponible. La falta de señal no deberá obligar al operador a desactivar controles de seguridad del dispositivo.")

heading(doc, "14", "CAMBIOS Y AUDITORÍA")
para(doc, "GE Control no modificará el modelo de captura puntual a rastreo continuo sin nueva versión y autorización contractual. El CLIENTE podrá solicitar información razonable sobre configuración y accesos conforme al Anexo de Tratamiento y Seguridad.")
para(doc, "Cada aceptación conservará versión, cliente, RFC, eventos habilitados, plazo, fecha y evidencia. Los cambios materiales deberán comunicarse a operadores antes de aplicarse.")

heading(doc, "APÉNDICE A", "CONFIGURACIÓN DE EVENTOS")
grid(doc, ["Evento", "Habilitado", "Ubicación requerida", "Alternativa"], [
    ("Inicio de viaje", "{{geo_inicio_habilitado}}", "{{geo_inicio_requerida}}", "{{geo_inicio_alternativa}}"),
    ("Llegada a origen", "{{geo_origen_habilitado}}", "{{geo_origen_requerida}}", "{{geo_origen_alternativa}}"),
    ("Carga concluida", "{{geo_carga_habilitado}}", "{{geo_carga_requerida}}", "{{geo_carga_alternativa}}"),
    ("Incidencia", "{{geo_incidencia_habilitado}}", "{{geo_incidencia_requerida}}", "{{geo_incidencia_alternativa}}"),
    ("Llegada a destino", "{{geo_destino_habilitado}}", "{{geo_destino_requerida}}", "{{geo_destino_alternativa}}"),
    ("Entrega concluida", "{{geo_entrega_habilitado}}", "{{geo_entrega_requerida}}", "{{geo_entrega_alternativa}}"),
], [2500, 1700, 2260, 2900], 7.8)

heading(doc, "APÉNDICE B", "CONSTANCIA DE INFORMACIÓN AL OPERADOR")
para(doc, "Esta constancia acredita recepción de información; no implica renuncia de derechos ni sustituye el aviso de privacidad del CLIENTE.")
grid(doc, ["Dato", "Registro"], [
    ("Operador", "{{operador_nombre}}"),
    ("Cliente / RFC", "{{cliente_nombre_legal}} · {{suscripcion_rfc}}"),
    ("Aviso y política entregados", "{{aviso_operador_version}} · {{politica_geo_version}}"),
    ("Medio y fecha", "{{medio_entrega}} · {{fecha_entrega}}"),
    ("Manifestación", "Declaro haber recibido información sobre eventos, finalidades, conservación, accesos, derechos y medio para reportar fallas."),
    ("Firma / evidencia", "{{operador_firma_o_evidencia}}"),
], [3300, 6060], 8.2)

heading(doc, "APÉNDICE C", "ACEPTACIÓN ENTRE LAS PARTES")
grid(doc, ["POR EL CLIENTE", "POR GE CONTROL"], [
    ("Nombre: {{cliente_representante}}\nCargo: {{cliente_cargo}}\nFirma: __________________________\nFecha: {{cliente_fecha_firma}}",
     "María José Mejía Ornelas\nGE Control\nFirma: __________________________\nFecha: {{gecontrol_fecha_firma}}"),
], [4680, 4680], 9.0)

heading(doc, "APÉNDICE INTERNO", "CAMPOS PARA SUPERADMIN", page_break_before=True)
para(doc, "Este apéndice se omite del PDF firmado. La configuración debe corresponder exactamente con la versión informada al CLIENTE y a sus operadores.", italic=True)
grid(doc, ["Grupo", "Campos y controles"], [
    ("Versión", "Anexo, aviso del operador, política del cliente, vigencia, RFC y eventos habilitados."),
    ("Captura", "Evento, latitud, longitud, precisión, fecha, viaje, operador, dispositivo y fuente."),
    ("Permiso", "Estado técnico, fecha, versión del Portal y evidencia; no usarlo como única base jurídica."),
    ("Accesos", "Roles autorizados, consulta, exportación, ticket, usuario y sello temporal."),
    ("Conservación", "Fecha de captura, vencimiento, hold por incidencia, motivo, liberación y eliminación."),
    ("Constancia", "Operador, documentos entregados, medio, fecha, firma o evidencia y revocación/cambio."),
], [2200, 7160], 8.2)

doc.core_properties.title = "Anexo Especial de Geolocalización del Operador - GE Control"
doc.core_properties.subject = "Portal del Operador - Captura puntual de ubicación"
doc.core_properties.author = "GE Control"
doc.core_properties.keywords = "geolocalización, operador, privacidad laboral, Portal Transporte, datos personales"
doc.save(DEST)
print(DEST)
