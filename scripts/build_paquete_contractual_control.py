from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LEGAL = ROOT / "output" / "legal"
PACKAGE = LEGAL / "Paquete_Contractual_GE_Control"
PACKAGE.mkdir(parents=True, exist_ok=True)

DOCS = [
    ("01", "Contrato Marco SaaS", "Contrato_Marco_SaaS_GE_Control_Plantilla"),
    ("02", "Orden de Servicio", "Orden_de_Servicio_Portal_Transporte_Plantilla"),
    ("03", "Acuerdo de Niveles de Servicio", "Acuerdo_Niveles_Servicio_SLA_Portal_Transporte"),
    ("04", "Aviso de Privacidad Integral", "Aviso_Privacidad_Integral_GE_Control"),
    ("05", "Anexo de Tratamiento y Seguridad", "Anexo_Tratamiento_Seguridad_Datos_GE_Control"),
    ("06", "Política de Soporte, Respaldos, Conservación y Exportación", "Politica_Soporte_Respaldos_Conservacion_Exportacion_GE_Control"),
    ("07", "Reglas de Consumo de Viajes Fiscales y Cancelaciones", "Reglas_Consumo_Viajes_Fiscales_Cancelaciones_GE_Control"),
    ("08", "Anexo Especial de Geolocalización del Operador", "Anexo_Especial_Geolocalizacion_Operador_GE_Control"),
]

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
LIGHT = "F6F1F2"
HEADER = "E9DDE0"
FONT = "Aptos"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_fields(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        xml = "\n".join(
            zf.read(name).decode("utf-8", "ignore")
            for name in zf.namelist()
            if name.endswith(".xml")
        )
    return sorted(set(re.findall(r"\{\{([A-Za-z0-9_]+)\}\}", xml)))


def category(name: str) -> str:
    if name.startswith(("contrato_", "orden_servicio_", "cotizacion_", "sla_", "dpa_", "politica_operativa_", "reglas_consumo_", "geo_version", "geo_fecha_", "aviso_operador_", "politica_geo_")):
        return "Documento y versión"
    if name.startswith("proveedor_") or name.startswith("gecontrol_") or name == "jurisdiccion_estado":
        return "Identidad legal de GE Control"
    if name.startswith("cliente_"):
        return "Cliente y firmantes"
    if name.startswith("suscripcion_") or name in {"subscription_reference", "tenant_reference"}:
        return "RFC, tenant y suscripción"
    if name.startswith(("plan_", "limite_", "viajes_", "periodicidad", "moneda", "subtotal", "iva", "total_", "condiciones_pago", "regla_renovacion", "fecha_activacion", "inicio_periodo_", "fin_periodo_", "legacy_", "grandfathered_", "initial_override")):
        return "Plan, precio y límites"
    if name.startswith(("portal_", "implementacion_", "otros_", "special_terms", "approved_notes")):
        return "Addons, descuentos y excepciones"
    if name.startswith("geo_") or name in {"operador_nombre", "operador_firma_o_evidencia", "fecha_entrega", "medio_entrega"}:
        return "Geolocalización y constancia"
    if name.startswith("region_") or name in {"dias_ventana_exportacion", "dias_retencion_respaldos"}:
        return "Tratamiento, infraestructura y conservación"
    if name.startswith("soporte_") or name in {"correo_soporte", "canal_emergencia"}:
        return "Soporte y SLA"
    if name.startswith("fecha_") or name.endswith("_fecha_firma"):
        return "Fechas y aceptación"
    return "Control administrativo"


def source(name: str, cat: str) -> str:
    if cat == "Identidad legal de GE Control":
        return "Perfil legal versionado de GE Control"
    if cat == "Cliente y firmantes":
        return "Expediente contractual del cliente"
    if cat == "RFC, tenant y suscripción":
        return "Cliente/RFC/suscripción resueltos por servidor"
    if cat == "Plan, precio y límites":
        return "Versión de plan y términos comerciales"
    if cat == "Addons, descuentos y excepciones":
        return "Cotización aprobada y autorización Superadmin"
    if cat == "Documento y versión":
        return "Motor documental/versiones publicadas"
    if cat == "Geolocalización y constancia":
        return "Configuración por RFC y evidencia del operador"
    if cat == "Tratamiento, infraestructura y conservación":
        return "Configuración jurídica/operativa aprobada"
    if cat == "Soporte y SLA":
        return "Política operativa y nivel contratado"
    if cat == "Fechas y aceptación":
        return "Evento de emisión, aceptación o firma"
    return "Superadmin con auditoría"


def control(name: str, cat: str) -> str:
    if cat == "Identidad legal de GE Control":
        return "Sólo configuración legal; cambio crea nueva versión"
    if name in {"tenant_reference", "subscription_reference", "suscripcion_rfc"}:
        return "Sólo servidor; nunca texto libre del navegador"
    if name.endswith("_version") or cat == "Documento y versión":
        return "Seleccionar versión publicada; no editar texto"
    if name.endswith(("_firma", "_firma_o_evidencia")) or "fecha_firma" in name or name == "fecha_aceptacion":
        return "Se genera por evento/evidencia; no editable después"
    if cat in {"Plan, precio y límites", "Addons, descuentos y excepciones"}:
        return "Editable en borrador; cambios sensibles requieren motivo"
    if cat == "Geolocalización y constancia":
        return "Editable por configuración autorizada; conservar evidencia"
    return "Editable en borrador según rol"


def freeze(name: str, cat: str) -> str:
    if cat == "Geolocalización y constancia" and name.startswith("geo_"):
        return "Al entregar aviso al operador"
    if name.endswith(("_firma", "_firma_o_evidencia")) or "fecha_firma" in name:
        return "Al registrar la firma/evidencia"
    if cat == "Documento y versión":
        return "Al emitir la versión"
    return "Al emitir cotización/orden; definitivo al aceptar"


def exposure(cat: str, name: str) -> str:
    if name in {"tenant_reference", "subscription_reference", "approved_notes", "initial_override"}:
        return "Interno; excluir del PDF del cliente"
    if cat == "Identidad legal de GE Control":
        return "Contrato privado; no página pública"
    if cat == "Geolocalización y constancia":
        return "Cliente/operador autorizado"
    return "Documento privado de las partes"


def required(name: str) -> bool:
    optional_tokens = (
        "portal_", "implementacion_", "otros_", "special_terms", "approved_notes",
        "initial_override", "legacy_flag", "grandfathered_flag", "canal_emergencia",
        "operador_", "fecha_entrega", "medio_entrega",
    )
    return not name.startswith(optional_tokens)


def validation(name: str) -> str:
    if name in {"cliente_rfc", "suscripcion_rfc", "proveedor_rfc"}:
        return "RFC normalizado y validado; mayúsculas; 12/13 caracteres"
    if "correo" in name or "contacto_privacidad" in name or "contacto_seguridad" in name:
        return "Correo válido"
    if name.startswith("fecha_") or name.endswith(("_fecha_vigencia", "_inicio", "_fin")):
        return "Fecha ISO y zona America/Mexico_City"
    if name in {"iva", "subtotal", "total_inicial", "total_recurrente"} or "precio" in name or "acordad" in name or "lista" in name:
        return "Decimal no negativo; MXN; cálculo del servidor"
    if name.startswith("limite_") or name in {"viajes_incluidos_mes", "dias_ventana_exportacion", "dias_retencion_respaldos"}:
        return "Entero no negativo; respetar regla del plan"
    if name.endswith(("_habilitado", "_requerida")):
        return "Booleano"
    if name.endswith("_version"):
        return "Referencia a versión existente e inmutable"
    return "Texto saneado y longitud limitada"


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, value):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(value))
    tc_w.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=16 if level == 1 else 13, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), size=10.3)


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for idx, label in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8.8, bold=True, color=WINE)
    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for row in rows:
        table_row = tbl.add_row()
        tr_pr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        cells = table_row.cells
        for idx, value in enumerate(row):
            set_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(str(value)), size=8.3)
    return tbl


all_fields: dict[str, set[str]] = {}
doc_fields: dict[str, list[str]] = {}
for number, title, stem in DOCS:
    source_docx = LEGAL / f"{stem}.docx"
    fields = extract_fields(source_docx)
    doc_fields[number] = fields
    for field in fields:
        all_fields.setdefault(field, set()).add(number)

records = []
for field in sorted(all_fields):
    cat = category(field)
    records.append({
        "key": field,
        "documents": sorted(all_fields[field]),
        "category": cat,
        "source": source(field, cat),
        "superadmin_control": control(field, cat),
        "required": required(field),
        "validation": validation(field),
        "freeze_at": freeze(field, cat),
        "exposure": exposure(cat, field),
    })

catalog = {
    "schema_version": "1.0",
    "generated_on": str(date.today()),
    "product": "GE Control - Portal Transporte",
    "rules": {
        "one_subscription_per_rfc": True,
        "additional_trip_packages_allowed": False,
        "pin_operators_unlimited": True,
        "operator_portal_trial_max_months": 3,
        "legal_text_direct_edit": False,
        "accepted_snapshots_immutable": True,
        "public_identity_uses_trade_name_only": True,
        "private_contract_uses_legal_identity": True,
    },
    "fields": records,
}
(PACKAGE / "superadmin_contract_fields.json").write_text(
    json.dumps(catalog, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
)
with (PACKAGE / "superadmin_contract_fields.csv").open("w", newline="", encoding="utf-8-sig") as fh:
    writer = csv.DictWriter(fh, fieldnames=list(records[0]))
    writer.writeheader()
    for record in records:
        row = dict(record)
        row["documents"] = ",".join(row["documents"])
        row["required"] = "sí" if row["required"] else "no"
        writer.writerow(row)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.82)
section.right_margin = Inches(.82)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)
normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for style_name, size in (("Heading 1", 16), ("Heading 2", 13)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 10)
    style.paragraph_format.space_after = Pt(7 if style_name == "Heading 1" else 5)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  CONTROL MAESTRO CONTRACTUAL"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Control maestro v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(32)
title.paragraph_format.space_after = Pt(4)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("PAQUETE CONTRACTUAL"), size=25, bold=True, color=WINE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(20)
set_font(sub.add_run("Control maestro de documentos y variables de Superadmin"), size=13, color=MUTED)

table(doc, ["Dato", "Valor"], [
    ("Producto", "GE Control - Portal Transporte"),
    ("Versión del paquete", "1.0"),
    ("Fecha de control", str(date.today())),
    ("Documentos jurídicos", "8"),
    ("Variables válidas", str(len(records))),
    ("Uso", "Plantilla interna; no sustituye los documentos firmados"),
], [2500, 6860])

heading(doc, "1. Arquitectura jurídica del paquete")
para(doc, "Los documentos no deben fusionarse en un solo contrato indiferenciado. Funcionan como un conjunto relacionado por folios y versiones; el Contrato Marco aporta las reglas comunes y la Orden de Servicio fija el acuerdo comercial de cada RFC.")
table(doc, ["Orden", "Documento", "Función"], [
    ("1", "Contrato Marco SaaS", "Relación jurídica general y reglas comunes."),
    ("2", "Orden de Servicio", "RFC, plan, límites, precio, addon y vigencia."),
    ("3", "SLA", "Disponibilidad, incidencias y atención."),
    ("4", "Aviso de Privacidad", "Información a titulares; documento público controlado."),
    ("5", "Tratamiento y Seguridad", "Responsable/encargado, medidas y subencargados."),
    ("6", "Política Operativa", "Soporte, respaldos, conservación y exportación."),
    ("7", "Reglas de Viajes", "Cómputo, Carta Porte, CFDI de ingreso y cancelaciones."),
    ("8", "Geolocalización", "Eventos puntuales, límites laborales y constancia."),
], [700, 3200, 5460])

heading(doc, "2. Capas de control")
table(doc, ["Capa", "Quién la controla", "Regla"], [
    ("Texto jurídico", "Administrador jurídico", "Sólo nuevas versiones; nunca edición directa por cliente."),
    ("Datos legales GE Control", "Superadmin autorizado", "Cambio exige evidencia y nueva versión del perfil legal."),
    ("Datos del cliente/RFC", "Superadmin", "Validados contra expediente; RFC aislado por suscripción."),
    ("Oferta comercial", "Superadmin comercial", "Editable sólo en borrador; descuentos con motivo y aprobador."),
    ("Configuración operativa", "Superadmin autorizado", "Debe corresponder a plan y anexos vigentes."),
    ("Snapshot emitido", "Sistema", "Inmutable; corrección mediante nueva versión."),
    ("Aceptación/firma", "Sistema y partes", "Sello temporal, actor, hash y evidencia."),
], [1900, 2300, 5160])

heading(doc, "3. Flujo obligatorio")
for item in [
    "Crear o seleccionar cliente contractual y RFC.",
    "Crear cotización en borrador desde una versión publicada de plan y precio.",
    "Revisar descuentos, addon, límites, vigencia e IVA.",
    "Emitir una versión inmutable de la cotización.",
    "Al aceptar, crear la suscripción por RFC y su Orden de Servicio versionada.",
    "Seleccionar las versiones jurídicas aplicables y resolver todas las variables requeridas.",
    "Generar PDF de cliente sin apéndices internos; validar que no queden marcadores.",
    "Registrar aceptación o firma, hash, archivo final y evidencia.",
    "Toda modificación posterior se formaliza con nueva versión, renovación o convenio; nunca se sobrescribe lo aceptado.",
]:
    bullet(doc, item)

heading(doc, "4. Reglas comerciales cerradas")
for item in [
    "Una suscripción corresponde a un RFC. Cada RFC adicional contrata otra suscripción.",
    "No se venden paquetes adicionales de viajes; al agotarse el límite se propone el plan siguiente.",
    "Los operadores PIN son ilimitados; administradores activos e invitados consumen cupo.",
    "El Portal del Operador es addon separado y su prueba no puede exceder tres meses.",
    "Un viaje fiscal se consume al timbrar la Carta Porte con UUID; su CFDI de ingreso relacionado no consume otro viaje.",
    "Los viajes no utilizados no se acumulan y una cancelación no devuelve automáticamente el consumo.",
    "Precios, límites, descuentos, IVA y cláusulas deben quedar congelados en el snapshot aceptado.",
]:
    bullet(doc, item)

heading(doc, "5. Estados y permisos mínimos")
table(doc, ["Objeto", "Estados mínimos", "Quién modifica"], [
    ("Cláusula/plan/precio", "draft · published · retired", "Superadmin autorizado; publicar exige auditoría."),
    ("Cotización", "draft · review · issued · accepted/closed · converted", "Comercial; aceptación y conversión por acciones separadas."),
    ("Orden de Servicio", "draft · issued · accepted · superseded", "Comercial/jurídico; emitida no se edita."),
    ("Suscripción", "planned · active · past_due · suspended · terminated", "Servidor y Superadmin con razón."),
    ("Documento firmado", "generated · sent · accepted/signed · superseded", "Sistema; archivo y hash inmutables."),
], [1900, 4660, 2800])

heading(doc, "6. Validaciones de generación")
for item in [
    "Bloquear generación si falta una variable obligatoria o aparece un marcador desconocido.",
    "Resolver tenant, RFC y subscription_id exclusivamente en servidor.",
    "Recalcular subtotal, descuento, IVA y total en servidor; no confiar en valores del navegador.",
    "Excluir anexos marcados como INTERNOS de cualquier PDF enviado o firmado.",
    "Conservar versión de plantilla, versión jurídica, JSON de datos, hash SHA-256 y PDF exacto.",
    "No mostrar el nombre personal del proveedor en páginas públicas; sí incluir identidad legal en contratos privados.",
    "Previsualizar y exigir aprobación antes de emitir, enviar, aceptar o activar.",
]:
    bullet(doc, item)

heading(doc, "7. Mapa por documento")
table(doc, ["No.", "Documento", "Campos", "Entrega"], [
    (number, title, len(doc_fields[number]), "Cliente" if number != "04" else "Público y cliente")
    for number, title, _ in DOCS
], [700, 5000, 1100, 2560])

heading(doc, "8. Catálogo exacto de variables")
para(doc, "La fuente técnica completa está en superadmin_contract_fields.json y superadmin_contract_fields.csv. La tabla siguiente es el inventario humano. “Docs” corresponde a la numeración del apartado anterior.")
rows = []
for rec in records:
    rows.append((
        rec["key"],
        ",".join(rec["documents"]),
        rec["category"],
        "Sí" if rec["required"] else "No",
        rec["freeze_at"],
    ))
first_chunk_size = 8
chunk_size = 10
chunks = [rows[:first_chunk_size]]
chunks.extend(
    rows[start:start + chunk_size]
    for start in range(first_chunk_size, len(rows), chunk_size)
)
for index, chunk in enumerate(chunks):
    if index:
        doc.add_page_break()
        p = para(doc, "CATÁLOGO DE VARIABLES · CONTINUACIÓN", after=7)
        set_font(p.runs[0], size=9, bold=True, color=MUTED)
    table(
        doc,
        ["Variable", "Docs", "Categoría", "Req.", "Congelación"],
        chunk,
        [2800, 700, 2450, 650, 2760],
    )

heading(doc, "9. Campos que no deben ser texto libre")
for item in [
    "RFC, tenant, suscripción, plan, versión, moneda, periodicidad y estados.",
    "Límites, precios, descuentos, IVA, totales y fechas de vigencia.",
    "Versiones jurídicas y folios de documentos relacionados.",
    "Permisos de geolocalización, eventos habilitados y plazos de conservación.",
    "Identidad legal de GE Control: sólo desde perfil legal versionado con evidencia.",
]:
    bullet(doc, item)

heading(doc, "10. Checklist antes de firma")
for item in [
    "Cliente, representante, RFC, domicilio y contacto contractual verificados.",
    "Plan, viajes, vehículos, administradores, precio, IVA y periodicidad coinciden con la cotización aceptada.",
    "Portal del Operador y cualquier promoción tienen inicio, fin y efecto al vencimiento.",
    "Todos los folios y versiones cruzadas coinciden.",
    "No existen marcadores {{...}}, notas internas ni campos técnicos visibles.",
    "PDF revisado visualmente; archivo, hash y snapshot almacenados.",
    "Firmas/aceptación y facultades del representante documentadas.",
]:
    bullet(doc, item)

guide_docx = PACKAGE / "00_Control_Maestro_Paquete_Contractual_GE_Control.docx"
doc.save(guide_docx)

manifest_files = []
for number, title, stem in DOCS:
    for ext in ("docx", "pdf"):
        src = LEGAL / f"{stem}.{ext}"
        dst = PACKAGE / f"{number}_{stem}.{ext}"
        shutil.copy2(src, dst)
        manifest_files.append({
            "number": number,
            "title": title,
            "format": ext,
            "file": dst.name,
            "sha256": sha256(dst),
            "fields": doc_fields[number] if ext == "docx" else [],
        })

manifest = {
    "package_version": "1.0",
    "generated_on": str(date.today()),
    "product": "GE Control - Portal Transporte",
    "documents": manifest_files,
    "guide": guide_docx.name,
    "field_catalog_json": "superadmin_contract_fields.json",
    "field_catalog_csv": "superadmin_contract_fields.csv",
}
(PACKAGE / "contract_package_manifest.json").write_text(
    json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
)

(PACKAGE / "LEEME.md").write_text(
    """# Paquete contractual GE Control

Este directorio contiene las ocho plantillas jurídicas, el control maestro y el catálogo exacto de variables para Superadmin.

## Regla de entrega

Los DOCX son plantillas de trabajo. Antes de entregar o firmar, el motor debe resolver todos los marcadores, excluir los apéndices internos y conservar un snapshot inmutable con hash del PDF final.

## Archivos técnicos

- `superadmin_contract_fields.json`: contrato de datos para implementación.
- `superadmin_contract_fields.csv`: matriz operativa revisable.
- `contract_package_manifest.json`: inventario y hashes.

Los borradores deben ser validados por abogado mexicano antes de su primera firma en producción.
""",
    encoding="utf-8",
)

print(guide_docx)
print(PACKAGE / "superadmin_contract_fields.json")
print(PACKAGE / "superadmin_contract_fields.csv")
