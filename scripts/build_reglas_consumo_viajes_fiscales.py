from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "legal"
OUT.mkdir(parents=True, exist_ok=True)
DEST = OUT / "Reglas_Consumo_Viajes_Fiscales_Cancelaciones_GE_Control.docx"

WINE = RGBColor(91, 15, 29)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(92, 103, 120)
FONT = "Aptos"
LIGHT = "F6F1F2"
HEADER = "E9DDE0"


def set_font(run, size=10.3, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def para(doc, text="", after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.14
    set_font(p.add_run(text), italic=italic)
    return p


def heading(doc, number, title, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    set_font(p.add_run(f"{number}. {title}"), size=14, bold=True, color=WINE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.42)
    p.paragraph_format.first_line_indent = Inches(-.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=9.8)


def kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        set_width(cells[0], 2800)
        set_width(cells[1], 6560)
        shade(cells[0], LIGHT)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=9.2, bold=True, color=WINE)
        set_font(cells[1].paragraphs[0].add_run(value), size=9.2)
    return table


def grid(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        set_width(cell, widths[idx])
        shade(cell, HEADER)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8.8, bold=True, color=WINE)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            set_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(value), size=font_size)
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.top_margin = Inches(.82)
section.bottom_margin = Inches(.82)
section.left_margin = Inches(.9)
section.right_margin = Inches(.9)
section.header_distance = Inches(.35)
section.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.3)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.14
for style_name, size in (("Heading 1", 14), ("Heading 2", 11.5)):
    style = doc.styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = WINE
    style.paragraph_format.space_before = Pt(11)
    style.paragraph_format.space_after = Pt(5)

head = section.header.paragraphs[0]
set_font(head.add_run("GE CONTROL  |  VIAJES FISCALES Y CANCELACIONES"), size=8.5, bold=True, color=MUTED)
foot = section.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("Reglas de consumo v1.0  ·  Página "), size=8, color=MUTED)
add_field(foot, "PAGE")
set_font(foot.add_run(" de "), size=8, color=MUTED)
add_field(foot, "NUMPAGES")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run("REGLAS DE CONSUMO DE VIAJES FISCALES,\nTIMBRADO, SUSTITUCIONES Y CANCELACIONES"), size=19.5, bold=True, color=WINE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(15)
set_font(subtitle.add_run("Portal Transporte · Anexo operativo del Contrato Marco SaaS"), size=12.5, color=MUTED)

kv_table(doc, [
    ("Versión", "{{reglas_consumo_version}}"),
    ("Vigencia", "{{reglas_consumo_fecha_vigencia}}"),
    ("Contrato Marco", "{{contrato_marco_folio}}"),
    ("Orden de Servicio", "{{orden_servicio_folio}}"),
    ("Cliente / RFC", "{{cliente_nombre_legal}} · {{suscripcion_rfc}}"),
    ("Plan y viajes incluidos", "{{plan_nombre}} · {{viajes_incluidos_mes}} viajes por ciclo"),
])

para(doc, "Estas Reglas definen la unidad comercial denominada Viaje Fiscal, cuándo se consume, cómo se relacionan sus comprobantes y qué sucede ante errores, sustituciones, cancelaciones o agotamiento del límite. No sustituyen la legislación fiscal ni determinan por sí mismas el tipo de CFDI que corresponde a cada operación.")

heading(doc, "1", "DEFINICIONES")
grid(doc, ["Término", "Definición contractual"], [
    ("Viaje Fiscal", "Unidad comercial de consumo asociada a una operación de transporte identificable dentro de un RFC y ciclo de suscripción."),
    ("Carta Porte primaria", "CFDI con Complemento Carta Porte o documento fiscal configurado como comprobante inicial del viaje, certificado y con UUID."),
    ("CFDI de ingreso relacionado", "Comprobante posterior vinculado al viaje y a su Carta Porte primaria conforme a la configuración contratada."),
    ("Timbrado exitoso", "Respuesta del PAC que genera UUID y sello digital del SAT."),
    ("Intento fallido", "Solicitud rechazada antes de obtener UUID."),
    ("Sustitución", "Emisión de un nuevo CFDI para corregir otro, con las relaciones y motivos exigibles."),
    ("Ciclo", "Periodo mensual de consumo indicado en la Orden de Servicio."),
], [2300, 7060])

heading(doc, "2", "UNIDAD COMERCIAL: VIAJE, NO TIMBRE")
para(doc, "La suscripción se comercializa por Viajes Fiscales incluidos, no por número aislado de timbres. En la configuración ordinaria de Portal Transporte, un Viaje Fiscal puede incluir:")
for item in [
    "una Carta Porte primaria timbrada; y",
    "un CFDI de ingreso posterior, correctamente relacionado con esa Carta Porte y generado dentro del mismo flujo.",
]:
    bullet(doc, item)
para(doc, "Aunque intervengan dos certificaciones, el CFDI de ingreso relacionado no consume un segundo Viaje. Los costos internos del PAC no alteran esta regla salvo que la Orden de Servicio identifique expresamente otro flujo fiscal.")

heading(doc, "3", "SECUENCIA CONFIGURADA")
para(doc, "Para el flujo ordinario contratado, Portal Transporte exige una Carta Porte primaria timbrada antes de habilitar el CFDI de ingreso relacionado. Un intento de CFDI de ingreso sin Carta Porte previa será rechazado y no consumirá un Viaje.")
para(doc, "Esta secuencia es una regla técnica y comercial de la configuración contratada. No significa que sea el único esquema permitido por la normativa fiscal. El CLIENTE deberá confirmar con sus asesores el tipo de CFDI, complemento, relación y orden que correspondan a su operación.")

heading(doc, "4", "MOMENTO DE CONSUMO")
grid(doc, ["Evento", "¿Consume viaje?", "Regla"], [
    ("Guardar borrador o validar datos", "No", "No existe certificación ni UUID."),
    ("Intento rechazado por Portal, PAC o SAT", "No", "Siempre que no se haya generado UUID."),
    ("Carta Porte primaria con UUID", "Sí", "El consumo se registra una sola vez en el ciclo y RFC correspondientes."),
    ("Descargar XML o PDF nuevamente", "No", "Es consulta del mismo comprobante."),
    ("CFDI de ingreso relacionado", "No adicional", "Debe pertenecer al mismo viaje y cumplir la secuencia configurada."),
    ("Cancelar un CFDI existente", "No adicional", "La cancelación no genera por sí misma otro consumo."),
    ("Emitir una sustitución con nuevo UUID", "Sí", "El nuevo comprobante fiscal origina consumo conforme a la sección 7."),
], [3300, 1800, 4260], 8.0)

heading(doc, "5", "ASIGNACIÓN AL RFC Y AL CICLO")
para(doc, "Cada Viaje se carga exclusivamente a la suscripción del RFC emisor utilizado en el timbrado. Los viajes de otro RFC, aun perteneciente al mismo cliente o grupo, no comparten bolsa salvo convenio Enterprise expresamente documentado.")
para(doc, "El consumo se asigna al ciclo vigente en la fecha y hora de certificación registrada por el PAC, bajo la zona America/Mexico_City. Cambiar posteriormente fechas operativas, rutas, referencias o el estado del viaje no mueve el consumo a otro ciclo.")

heading(doc, "6", "CANCELACIONES")
para(doc, "Cancelar una Carta Porte o CFDI no restituye automáticamente el Viaje consumido, porque existió una certificación, trazabilidad, almacenamiento, procesamiento y costo del proveedor.")
para(doc, "El CLIENTE deberá seleccionar el motivo de cancelación correcto, proporcionar el folio sustituto cuando sea exigible y gestionar la aceptación del receptor cuando corresponda. El estado puede permanecer solicitado, en proceso, rechazado, vigente o cancelado según SAT y PAC.")
grid(doc, ["Situación", "Tratamiento comercial"], [
    ("Operación no realizada", "La cancelación no devuelve el Viaje ya consumido."),
    ("Error de captura del CLIENTE", "La cancelación no devuelve el Viaje; una nueva emisión consume otro."),
    ("Receptor niega o no concluye la cancelación", "No modifica el consumo mientras el UUID original exista."),
    ("Cancelación sin sustitución", "No genera consumo adicional, pero tampoco reintegro."),
    ("Cancelación solicitada durante suspensión de pago", "Se procurará habilitarla cuando sea técnica y legalmente posible."),
], [3500, 5860], 8.2)

heading(doc, "7", "SUSTITUCIONES Y CORRECCIONES")
para(doc, "Cuando la operación subsista y deba corregirse un CFDI, el CLIENTE deberá seguir el procedimiento, motivo de cancelación y relación aplicables. La generación de un comprobante sustituto con UUID consume un nuevo Viaje, aunque posteriormente se cancele el original.")
para(doc, "Para efectos comerciales, una sustitución es una nueva certificación. Portal Transporte deberá conservar UUID original, UUID sustituto, tipo de relación, motivo, usuario, fecha, RFC y resultado de cancelación.")
para(doc, "La clave 04 de relación entre CFDI y el motivo 01 de cancelación cumplen funciones distintas dentro del esquema fiscal; el CLIENTE no deberá tratarlos como equivalentes. La plataforma podrá orientar o validar campos, pero no sustituye el criterio fiscal del emisor.")

heading(doc, "8", "AJUSTE COMPENSATORIO POR ERROR DE GE CONTROL")
para(doc, "GE Control podrá otorgar un ajuste compensatorio cuando una nueva certificación haya sido necesaria exclusivamente por un defecto reproducible e imputable a la plataforma, y no por datos, configuración, decisión, omisión, conectividad o instrucción del CLIENTE.")
for item in [
    "La solicitud deberá presentarse dentro de diez días hábiles desde el comprobante sustituto.",
    "Debe identificar RFC, viaje, UUID original, UUID sustituto y ticket del incidente.",
    "GE Control verificará registros técnicos, causa y ausencia de modificación relevante del CLIENTE.",
    "El ajuste se reflejará como movimiento de auditoría; no borra UUID ni altera documentos fiscales.",
]:
    bullet(doc, item)
para(doc, "El ajuste no constituye reembolso en efectivo ni reconocimiento de responsabilidad. Su finalidad es neutralizar un consumo comercial cuando la evidencia confirme un error interno.")

heading(doc, "9", "AGOTAMIENTO DEL LÍMITE")
para(doc, "Al alcanzar los Viajes incluidos, Portal Transporte podrá impedir nuevas Cartas Porte primarias y nuevos Viajes. Continuarán, cuando sea técnicamente posible, consulta, descarga, cancelación y generación del CFDI de ingreso ya relacionado con un Viaje consumido.")
para(doc, "GE Control no venderá paquetes aislados de viajes adicionales bajo los planes estándar. Para aumentar capacidad, el CLIENTE deberá solicitar el cambio al plan siguiente. El upgrade y su fecha efectiva se documentarán en una Orden de Servicio o modificación.")
para(doc, "Las solicitudes concurrentes se evaluarán contra el saldo disponible. Si sólo queda un Viaje, la primera certificación exitosa lo consume y las demás deberán rechazarse antes de obtener UUID.")

heading(doc, "10", "RENOVACIÓN, ACUMULACIÓN Y CAMBIO DE PLAN")
para(doc, "Los Viajes no utilizados vencen al cerrar el ciclo y no se acumulan, transfieren, convierten en saldo, reembolsan ni compensan con otros RFC.")
grid(doc, ["Evento", "Efecto"], [
    ("Inicio de nuevo ciclo", "Se habilita la cantidad incluida en el plan vigente."),
    ("Upgrade durante el ciclo", "Aplicará en la fecha y con el ajuste de precio indicados en el documento modificatorio."),
    ("Downgrade", "Aplicará normalmente al siguiente ciclo y sólo si el consumo y vehículos cumplen el nuevo límite."),
    ("Terminación", "Los Viajes disponibles expiran en la fecha efectiva; subsisten exportación y obligaciones fiscales."),
    ("Plan legado", "Se respeta su límite documentado hasta que sea sustituido mediante aceptación expresa."),
], [3000, 6360], 8.3)

heading(doc, "11", "MEDICIÓN, TABLERO Y EVIDENCIA")
para(doc, "El tablero mostrará, cuando la función esté disponible, Viajes incluidos, consumidos, disponibles y movimientos del ciclo. El registro autoritativo será la bitácora de certificaciones y ajustes de GE Control, conciliada con respuestas del PAC.")
grid(doc, ["Campo mínimo", "Finalidad"], [
    ("subscription_id, RFC y ciclo", "Determinar la suscripción que soporta el consumo."),
    ("viaje_id y tipo de evento", "Relacionar operación y movimiento."),
    ("UUID, fecha y PAC", "Acreditar certificación exitosa."),
    ("Documento y relación", "Distinguir Carta Porte, ingreso, cancelación o sustitución."),
    ("Movimiento", "Consumo, ajuste compensatorio o corrección de conciliación."),
    ("Actor y evidencia", "Usuario, sistema, ticket, motivo y sello temporal."),
], [3100, 6260], 8.3)
para(doc, "Una demora de actualización visual no cambia el consumo real. GE Control corregirá diferencias de presentación o conciliación sin modificar XML, UUID ni estatus fiscal.")

heading(doc, "12", "ACLARACIONES")
para(doc, "El CLIENTE podrá objetar un movimiento dentro de diez días hábiles desde que aparezca en el tablero o reporte. Deberá aportar los identificadores y explicar la inconsistencia.")
para(doc, "GE Control responderá con evidencia razonable dentro de cinco días hábiles, salvo investigación compleja o dependencia del PAC. La objeción no suspende automáticamente límites, facturación ni obligaciones fiscales.")

heading(doc, "13", "RESPONSABILIDAD FISCAL DEL CLIENTE")
for item in [
    "Determinar si corresponde CFDI de ingreso, traslado, complemento Carta Porte u otro esquema.",
    "Proporcionar datos completos, vigentes y correctos y revisar el borrador antes de timbrar.",
    "Custodiar certificados, contraseñas y autorizaciones.",
    "Verificar relaciones, motivos, plazos y aceptación de cancelaciones.",
    "Consultar a sus asesores y conservar los documentos exigibles.",
]:
    bullet(doc, item)
para(doc, "GE Control proporciona una herramienta tecnológica y controles configurables; no actúa como contador, asesor fiscal, transportista, PAC ni autoridad. Una validación técnica exitosa no garantiza la deducibilidad, acreditamiento o cumplimiento integral de la operación.")

heading(doc, "14", "CAMBIOS Y VERSIONADO")
para(doc, "Estas Reglas sólo podrán modificarse prospectivamente. Los conceptos de Viaje, momento de consumo, no acumulación, límites y tratamiento de sustituciones aplicables a una Orden de Servicio quedarán vinculados a su versión aceptada.")
para(doc, "Los cambios derivados de normas del SAT o del PAC podrán implementarse para mantener operación y cumplimiento. Si alteran materialmente la unidad comercial o precio, requerirán aviso y documento modificatorio.")

heading(doc, "APÉNDICE A", "EJEMPLOS DE CÓMPUTO")
grid(doc, ["Caso", "Resultado"], [
    ("Carta Porte timbrada + CFDI de ingreso relacionado", "1 Viaje total."),
    ("Tres intentos fallidos sin UUID y uno exitoso", "1 Viaje total."),
    ("Carta Porte timbrada y después cancelada sin sustitución", "1 Viaje; no se reintegra."),
    ("Carta Porte errónea y nueva Carta Porte sustituta con UUID", "2 Viajes, salvo ajuste aprobado por error de GE Control."),
    ("Reimpresión o descarga múltiple del mismo UUID", "0 Viajes adicionales."),
    ("Ingreso solicitado sin Carta Porte previa en el flujo configurado", "Se rechaza; 0 Viajes mientras no exista UUID."),
    ("Ingreso relacionado con viaje ya consumido", "0 Viajes adicionales."),
], [6000, 3360], 8.4)

heading(doc, "APÉNDICE B", "ACEPTACIÓN")
para(doc, "Estas Reglas se aceptan como parte de la Orden de Servicio. La firma electrónica, aceptación verificable o firma autógrafa producirá los efectos previstos en el Contrato Marco.")
grid(doc, ["POR EL CLIENTE", "POR GE CONTROL"], [
    ("Nombre: {{cliente_representante}}\nCargo: {{cliente_cargo}}\nFirma: __________________________\nFecha: {{cliente_fecha_firma}}",
     "María José Mejía Ornelas\nGE Control\nFirma: __________________________\nFecha: {{gecontrol_fecha_firma}}"),
], [4680, 4680], 9.0)

heading(doc, "APÉNDICE INTERNO", "CAMPOS PARA SUPERADMIN", page_break_before=True)
para(doc, "Este apéndice se omite del PDF firmado. Los movimientos fiscales son inmutables; Superadmin sólo administra configuración versionada y ajustes compensatorios auditados.", italic=True)
grid(doc, ["Grupo", "Campos y controles"], [
    ("Configuración", "Versión, ciclo, RFC, subscription_id, plan, viajes incluidos y zona horaria."),
    ("Consumo", "Viaje, UUID, fecha PAC, tipo de documento, movimiento, saldo anterior y posterior."),
    ("Relaciones", "UUID Carta Porte, UUID ingreso, UUID sustituto, tipo de relación y motivo de cancelación."),
    ("Ajustes", "Solicitud, ticket, evidencia, causa, autorización, cantidad, usuario y sello temporal."),
    ("Límites", "Alertas, agotamiento, bloqueo, solicitud de upgrade y fecha efectiva."),
    ("Conciliación", "Estado Portal/PAC/SAT, diferencia detectada, resolución y evidencia, sin editar XML o UUID."),
], [2200, 7160], 8.2)

doc.core_properties.title = "Reglas de Consumo de Viajes Fiscales y Cancelaciones - GE Control"
doc.core_properties.subject = "Portal Transporte - Reglas comerciales y operativas de consumo"
doc.core_properties.author = "GE Control"
doc.core_properties.keywords = "viajes fiscales, Carta Porte, CFDI, timbrado, cancelación, sustitución"
doc.save(DEST)
print(DEST)
